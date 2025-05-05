# #!/usr/bin/env python3
# import importlib  
# import pkgutil
# import random
# import json

# # 1) LIST all your generator modules (sans “.py”)
# GENERATORS = [
#     "generate_email",
#     "generate_upi_id",
#     "generate_phone_no",
#     "generate_ifsc_code",
#     "generate_bank_acc_no",
#     "generate_pan",
#     "generate_date_and_time",
#     "generate_url",
#     "generate_ip_address",
#     "generate_ration_card",
#     "generate_address",
#     "generate_credit_card",
#     "generate_debit_card",
#     "generate_gender",
#     "generate_gst_no",
#     "generate_imei",
#     "generate_mac_address",
#     "generate_dl",
#     "generate_dob",
#     "generate_passport",
#     "generate_voter_id",
#     "generate_vehicle_reg_no",
# ]
# # 2) Dynamically import each module and call its generator
# def collect_sentences(count_per_module=10):
#     all_texts = []
#     for mod_name in GENERATORS:
#         mod = importlib.import_module(mod_name)
#         # assume each module exports a function named generate_XXX_variations
#         fn = next(
#             getattr(mod, name)
#             for name in dir(mod)
#             if name.startswith("generate_") and callable(getattr(mod, name))
#         )
#         # collect a few examples from each PII type
#         records = fn(count=count_per_module)
#         all_texts += [rec["text"] for rec in records]
#     return all_texts

# # 3) Build paragraphs
# def build_paragraphs(texts, num_paras, sents_per_para):
#     paras = []
#     for _ in range(num_paras):
#         paras.append(" ".join(random.sample(texts, sents_per_para)))
#     return paras

# # 4) Write to file
# def write_txt(paragraphs, filename):
#     with open(filename, "w", encoding="utf-8") as f:
#         for p in paragraphs:
#             f.write(p + "\n\n")
#     print(f"→ {filename} with {len(paragraphs)} paragraphs written")

# if __name__ == "__main__":
#     # you can tweak these
#     PARAS = 10
#     SENTS = 4
#     pool = collect_sentences(count_per_module=PARAS)
#     # for each category of file you asked for:
#     for category, out_name in [
#         ("structured", "structured.txt"),
#         ("compressed", "compressed.txt"),
#         ("unstructured", "unstructured.txt"),
#     ]:
#         # 10 paragraphs each
#         paras = build_paragraphs(pool, num_paras=10, sents_per_para=SENTS)
#         write_txt(paras, out_name)





# #!/usr/bin/env python3
# import importlib
# import os
# import sys
# import random
# import json
# import pandas as pd

# # these two imports are for PDF and DOCX output
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# from docx import Document

# # -----------------------------------------------------------------------------
# # 1. Configuration
# # -----------------------------------------------------------------------------
# # how many sentences (rows) per module to pull
# DEFAULT_COUNT = 50

# # your file‐type buckets
# STRUCTURED = ["csv", "xlsx", "xls", "parquet", "json"]
# COMPRESSED = ["pst"]
# UNSTRUCTURED = ["txt", "log", "html", "pdf", "docx", "sql", "diff"]

# # where to write
# OUT_DIR = "mixed_output"
# os.makedirs(OUT_DIR, exist_ok=True)

# # -----------------------------------------------------------------------------
# # 2. Discover and import your generators
# #    assumes each generate_*.py lives next to this script
# # -----------------------------------------------------------------------------
# MODULES = []
# for fn in os.listdir("."):
#     if fn.startswith("generate_") and fn.endswith(".py"):
#         mod_name = fn[:-3]
#         MODULES.append(importlib.import_module(mod_name))

# # helper to find each module's generation function
# def get_generator(mod):
#     # e.g. module generate_email has generate_email_variations
#     base = mod.__name__.split("generate_")[-1]
#     fname = f"generate_{base}_variations"
#     return getattr(mod, fname)

# # -----------------------------------------------------------------------------
# # 3. Collect a giant mixed pool of sentences
# # -----------------------------------------------------------------------------
# def collect_records(count_per_module):
#     pool = []
#     for mod in MODULES:
#         gen = get_generator(mod)
#         pool.extend(gen(count_per_module))
#     random.shuffle(pool)
#     return pool

# # -----------------------------------------------------------------------------
# # 4. Emit structured files via pandas
# # -----------------------------------------------------------------------------
# def emit_structured(df: pd.DataFrame):
#     # CSV
#     df.to_csv(os.path.join(OUT_DIR, "output.csv"), index=False)
#     # JSON
#     df.to_json(os.path.join(OUT_DIR, "output.json"), orient="records", indent=2)
#     # Excel .xlsx
#     df.to_excel(os.path.join(OUT_DIR, "output.xlsx"), index=False)
#     # Excel .xls
#     df.to_excel(os.path.join(OUT_DIR, "output.xls"), index=False, engine="xlwt")
#     # Parquet
#     df.to_parquet(os.path.join(OUT_DIR, "output.parquet"), index=False)

# # -----------------------------------------------------------------------------
# # 5. Emit compressed / “.pst” as just a JSON dump with .pst extension
# # -----------------------------------------------------------------------------
# def emit_compressed(records):
#     with open(os.path.join(OUT_DIR, "output.pst"), "w", encoding="utf-8") as f:
#         json.dump(records, f, ensure_ascii=False, indent=2)

# # -----------------------------------------------------------------------------
# # 6. Emit unstructured formats
# # -----------------------------------------------------------------------------
# def emit_txt(records):
#     with open(os.path.join(OUT_DIR, "output.txt"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(r["text"] + "\n\n")

# def emit_log(records):
#     with open(os.path.join(OUT_DIR, "output.log"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(f"{r['variation'].upper()}: {r['text']}\n")

# def emit_html(records):
#     with open(os.path.join(OUT_DIR, "output.html"), "w", encoding="utf-8") as f:
#         f.write("<!doctype html><html><body>\n")
#         for r in records:
#             f.write(f"<p>{r['text']}</p>\n")
#         f.write("</body></html>")

# def emit_pdf(records):
#     c = canvas.Canvas(os.path.join(OUT_DIR, "output.pdf"), pagesize=letter)
#     width, height = letter
#     y = height - 40
#     for r in records:
#         text = r["text"]
#         c.drawString(40, y, text)
#         y -= 14
#         if y < 40:
#             c.showPage()
#             y = height - 40
#     c.save()

# def emit_docx(records):
#     doc = Document()
#     for r in records:
#         doc.add_paragraph(r["text"])
#     doc.save(os.path.join(OUT_DIR, "output.docx"))

# def emit_sql(records):
#     # assume a table 'sentences(text, variation, is_valid)'
#     with open(os.path.join(OUT_DIR, "output.sql"), "w", encoding="utf-8") as f:
#         for r in records:
#             txt = r["text"].replace("'", "''")
#             var = r["variation"]
#             val = '1' if r.get("is_valid") else '0'
#             f.write(f"INSERT INTO sentences(text, variation, is_valid) VALUES('{txt}', '{var}', {val});\n")

# def emit_diff(records):
#     with open(os.path.join(OUT_DIR, "output.diff"), "w", encoding="utf-8") as f:
#         f.write("--- original\n+++ mixed\n")
#         for r in records:
#             f.write(f"+ {r['text']}\n")
#         f.write("\n")

# # -----------------------------------------------------------------------------
# # 7. Main entrypoint
# # -----------------------------------------------------------------------------
# if __name__ == "__main__":
#     count = int(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_COUNT
#     recs = collect_records(count)
#     df   = pd.DataFrame(recs)

#     # structured
#     emit_structured(df)

#     # compressed
#     emit_compressed(recs)

#     # unstructured
#     emit_txt(recs)
#     emit_log(recs)
#     emit_html(recs)
#     emit_pdf(recs)
#     emit_docx(recs)
#     emit_sql(recs)
#     emit_diff(recs)

#     print(f"✅ Generated {count} records each from {len(MODULES)} modules")
#     print(f"   → outputs in folder: {OUT_DIR}")






# #!/usr/bin/env python3
# import importlib
# import os
# import sys
# import random
# import json
# import pandas as pd

# # these two imports are for PDF and DOCX output
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# from docx import Document

# # -----------------------------------------------------------------------------
# # 1. Configuration
# # -----------------------------------------------------------------------------
# # how many sentences (rows) per module to pull
# DEFAULT_COUNT = 50

# # your file‐type buckets
# STRUCTURED   = ["csv", "xlsx", "xls", "parquet", "json"]
# COMPRESSED   = ["pst"]
# UNSTRUCTURED = ["txt", "log", "html", "pdf", "docx", "sql", "diff"]

# # where to write
# OUT_DIR = "mixed_output"
# os.makedirs(OUT_DIR, exist_ok=True)

# # -----------------------------------------------------------------------------
# # 2. Discover and import your generators
# #    assumes each generate_*.py lives next to this script
# # -----------------------------------------------------------------------------
# MODULES = []
# for fn in os.listdir("."):
#     if fn.startswith("generate_") and fn.endswith(".py"):
#         mod_name = fn[:-3]
#         MODULES.append(importlib.import_module(mod_name))

# # -----------------------------------------------------------------------------
# # 3. Helper to find each module's generation function
# #    now scans for any generate_*_variations callable
# # -----------------------------------------------------------------------------
# def get_generator(mod):
#     for attr in dir(mod):
#         if attr.startswith("generate_") and attr.endswith("_variations"):
#             fn = getattr(mod, attr)
#             if callable(fn):
#                 return fn
#     raise AttributeError(
#         f"No generate_*_variations() function found in module {mod.__name__}"
#     )

# # -----------------------------------------------------------------------------
# # 4. Collect a giant mixed pool of sentences
# # -----------------------------------------------------------------------------
# def collect_records(count_per_module):
#     pool = []
#     for mod in MODULES:
#         gen = get_generator(mod)
#         pool.extend(gen(count_per_module))
#     random.shuffle(pool)
#     return pool

# # -----------------------------------------------------------------------------
# # 5. Emit structured files via pandas
# # -----------------------------------------------------------------------------
# def emit_structured(df: pd.DataFrame):
#     # CSV
#     df.to_csv(os.path.join(OUT_DIR, "output.csv"), index=False)
#     # JSON
#     df.to_json(os.path.join(OUT_DIR, "output.json"), orient="records", indent=2)
#     # Excel .xlsx
#     df.to_excel(os.path.join(OUT_DIR, "output.xlsx"), index=False)
#     # Excel .xls
#     df.to_excel(os.path.join(OUT_DIR, "output.xls"), index=False, engine="xlwt")
#     # Parquet
#     df.to_parquet(os.path.join(OUT_DIR, "output.parquet"), index=False)

# # -----------------------------------------------------------------------------
# # 6. Emit compressed / “.pst” as just a JSON dump with .pst extension
# # -----------------------------------------------------------------------------
# def emit_compressed(records):
#     with open(os.path.join(OUT_DIR, "output.pst"), "w", encoding="utf-8") as f:
#         json.dump(records, f, ensure_ascii=False, indent=2)

# # -----------------------------------------------------------------------------
# # 7. Emit unstructured formats
# # -----------------------------------------------------------------------------
# def emit_txt(records):
#     with open(os.path.join(OUT_DIR, "output.txt"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(r["text"] + "\n\n")

# def emit_log(records):
#     with open(os.path.join(OUT_DIR, "output.log"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(f"{r['variation'].upper()}: {r['text']}\n")

# def emit_html(records):
#     with open(os.path.join(OUT_DIR, "output.html"), "w", encoding="utf-8") as f:
#         f.write("<!doctype html><html><body>\n")
#         for r in records:
#             f.write(f"<p>{r['text']}</p>\n")
#         f.write("</body></html>")

# def emit_pdf(records):
#     c = canvas.Canvas(os.path.join(OUT_DIR, "output.pdf"), pagesize=letter)
#     width, height = letter
#     y = height - 40
#     for r in records:
#         text = r["text"]
#         c.drawString(40, y, text)
#         y -= 14
#         if y < 40:
#             c.showPage()
#             y = height - 40
#     c.save()

# def emit_docx(records):
#     doc = Document()
#     for r in records:
#         doc.add_paragraph(r["text"])
#     doc.save(os.path.join(OUT_DIR, "output.docx"))

# def emit_sql(records):
#     with open(os.path.join(OUT_DIR, "output.sql"), "w", encoding="utf-8") as f:
#         for r in records:
#             txt = r["text"].replace("'", "''")
#             var = r["variation"]
#             val = '1' if r.get("is_valid") else '0'
#             f.write(
#                 f"INSERT INTO sentences(text, variation, is_valid) "
#                 f"VALUES('{txt}', '{var}', {val})\n"
#             )

# def emit_diff(records):
#     with open(os.path.join(OUT_DIR, "output.diff"), "w", encoding="utf-8") as f:
#         f.write("--- original\n+++ mixed\n")
#         for r in records:
#             f.write(f"+ {r['text']}\n")
#         f.write("\n")

# # -----------------------------------------------------------------------------
# # 8. Main entrypoint
# # -----------------------------------------------------------------------------
# if __name__ == "__main__":
#     count = int(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_COUNT
#     recs  = collect_records(count)
#     df    = pd.DataFrame(recs)

#     # structured
#     emit_structured(df)

#     # compressed
#     emit_compressed(recs)

#     # unstructured
#     emit_txt(recs)
#     emit_log(recs)
#     emit_html(recs)
#     emit_pdf(recs)
#     emit_docx(recs)
#     emit_sql(recs)
#     emit_diff(recs)

#     print(f"✅ Generated {count} records each from {len(MODULES)} modules")
#     print(f"   → outputs in folder: {OUT_DIR}")



# #!/usr/bin/env python3
# import importlib
# import os
# import sys
# import random
# import json
# import pandas as pd

# # these two imports are for PDF and DOCX output
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# from docx import Document

# # -----------------------------------------------------------------------------
# # 1. Configuration
# # -----------------------------------------------------------------------------
# DEFAULT_COUNT = 50

# STRUCTURED   = ["csv", "xlsx", "xls", "parquet", "json"]
# COMPRESSED   = ["pst"]
# UNSTRUCTURED = ["txt", "log", "html", "pdf", "docx", "sql", "diff"]

# OUT_DIR = "mixed_output"
# os.makedirs(OUT_DIR, exist_ok=True)

# # -----------------------------------------------------------------------------
# # 2. Discover and import your generators
# # -----------------------------------------------------------------------------
# MODULES = []
# for fn in os.listdir("."):
#     if fn.startswith("generate_") and fn.endswith(".py"):
#         MODULES.append(importlib.import_module(fn[:-3]))

# # -----------------------------------------------------------------------------
# # 3. Helper to find each module's generation function
# # -----------------------------------------------------------------------------
# def get_generator(mod):
#     for name in dir(mod):
#         if name.endswith("_variations"):
#             fn = getattr(mod, name)
#             if callable(fn):
#                 return fn
#     raise AttributeError(f"No '*_variations' function found in {mod.__name__}")

# # -----------------------------------------------------------------------------
# # 4. Collect a giant mixed pool of records
# # -----------------------------------------------------------------------------
# def collect_records(count_per_module):
#     pool = []
#     for mod in MODULES:
#         gen = get_generator(mod)
#         pool.extend(gen(count_per_module))
#     random.shuffle(pool)
#     return pool

# # -----------------------------------------------------------------------------
# # 5. Emit structured files via pandas
# # -----------------------------------------------------------------------------
# def emit_structured(df: pd.DataFrame):
#     # CSV
#     df.to_csv(os.path.join(OUT_DIR, "output.csv"), index=False)
#     # JSON
#     df.to_json(os.path.join(OUT_DIR, "output.json"), orient="records", indent=2)
#     # XLSX (requires openpyxl)
#     df.to_excel(os.path.join(OUT_DIR, "output.xlsx"), index=False)
#     # Parquet (optional)
#     try:
#         df.to_parquet(os.path.join(OUT_DIR, "output.parquet"), index=False)
#     except ImportError as e:
#         print("⚠️  Skipping Parquet output – install pyarrow or fastparquet to enable it")

# # -----------------------------------------------------------------------------
# # 6. Emit compressed / “.pst” as just a JSON dump
# # -----------------------------------------------------------------------------
# def emit_compressed(records):
#     with open(os.path.join(OUT_DIR, "output.pst"), "w", encoding="utf-8") as f:
#         json.dump(records, f, ensure_ascii=False, indent=2)

# # -----------------------------------------------------------------------------
# # 7. Emit unstructured formats
# # -----------------------------------------------------------------------------
# def emit_txt(records):
#     with open(os.path.join(OUT_DIR, "output.txt"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(r["text"] + "\n\n")

# def emit_log(records):
#     with open(os.path.join(OUT_DIR, "output.log"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(f"{r['variation'].upper()}: {r['text']}\n")

# def emit_html(records):
#     with open(os.path.join(OUT_DIR, "output.html"), "w", encoding="utf-8") as f:
#         f.write("<!doctype html><html><body>\n")
#         for r in records:
#             f.write(f"<p>{r['text']}</p>\n")
#         f.write("</body></html>")

# def emit_pdf(records):
#     c = canvas.Canvas(os.path.join(OUT_DIR, "output.pdf"), pagesize=letter)
#     width, height = letter
#     y = height - 40
#     for r in records:
#         c.drawString(40, y, r["text"])
#         y -= 14
#         if y < 40:
#             c.showPage()
#             y = height - 40
#     c.save()

# def emit_docx(records):
#     doc = Document()
#     for r in records:
#         doc.add_paragraph(r["text"])
#     doc.save(os.path.join(OUT_DIR, "output.docx"))

# def emit_sql(records):
#     with open(os.path.join(OUT_DIR, "output.sql"), "w", encoding="utf-8") as f:
#         for r in records:
#             txt = r["text"].replace("'", "''")
#             val = '1' if r.get("is_valid") else '0'
#             f.write(f"INSERT INTO sentences(text, variation, is_valid) VALUES('{txt}', '{r['variation']}', {val});\n")

# def emit_diff(records):
#     with open(os.path.join(OUT_DIR, "output.diff"), "w", encoding="utf-8") as f:
#         f.write("--- original\n+++ mixed\n")
#         for r in records:
#             f.write(f"+ {r['text']}\n")
#         f.write("\n")

# # -----------------------------------------------------------------------------
# # 8. Main entrypoint
# # -----------------------------------------------------------------------------
# if __name__ == "__main__":
#     count = int(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_COUNT
#     recs  = collect_records(count)
#     df    = pd.DataFrame(recs)

#     # structured
#     emit_structured(df)
#     # compressed
#     emit_compressed(recs)
#     # unstructured
#     emit_txt(recs)
#     emit_log(recs)
#     emit_html(recs)
#     emit_pdf(recs)
#     emit_docx(recs)
#     emit_sql(recs)
#     emit_diff(recs)

#     print(f"✅ Generated {count} records each from {len(MODULES)} modules")
#     print(f"   → outputs in folder: {OUT_DIR}")




# #!/usr/bin/env python3
# import importlib
# import os
# import sys
# import random
# import json
# import pandas as pd

# # these two imports are for PDF and DOCX output
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# from docx import Document

# # -----------------------------------------------------------------------------
# # 1. Configuration
# # -----------------------------------------------------------------------------
# DEFAULT_COUNT   = 50    # how many sentences per module to pull when no CLI arg
# PARA_COUNT      = 5     # how many paragraphs to write
# SENTS_PER_PARA  = 5     # how many sentences per paragraph

# STRUCTURED      = ["csv", "xlsx", "xls", "parquet", "json"]
# COMPRESSED      = ["pst"]
# UNSTRUCTURED    = ["txt", "log", "html", "pdf", "docx", "sql", "diff"]

# OUT_DIR = "mixed_output"
# os.makedirs(OUT_DIR, exist_ok=True)

# # -----------------------------------------------------------------------------
# # 2. Discover and import your generators
# # -----------------------------------------------------------------------------
# MODULES = []
# for fn in os.listdir("."):
#     if fn.startswith("generate_") and fn.endswith(".py"):
#         mod_name = fn[:-3]
#         MODULES.append(importlib.import_module(mod_name))

# def get_generator(mod):
#     """
#     Given module generate_<X>.py, return its generate_<X>_variations function.
#     Some modules name theirs slightly differently, so we try a couple of patterns.
#     """
#     base = mod.__name__.split("generate_")[-1]
#     candidates = [
#         f"generate_{base}_variations",
#         f"generate_{base}s",               # fallback plural
#         f"generate_{base}",                # fallback
#     ]
#     for name in candidates:
#         if hasattr(mod, name):
#             return getattr(mod, name)
#     raise AttributeError(f"No generator function found in {mod.__name__}")

# # -----------------------------------------------------------------------------
# # 3. Collect a giant mixed pool of sentences
# # -----------------------------------------------------------------------------
# def collect_records(count_per_module):
#     pool = []
#     for mod in MODULES:
#         gen = get_generator(mod)
#         pool.extend(gen(count_per_module))
#     random.shuffle(pool)
#     return pool

# # -----------------------------------------------------------------------------
# # 4. Emit structured files via pandas
# # -----------------------------------------------------------------------------
# def emit_structured(df: pd.DataFrame):
#     df.to_csv(os.path.join(OUT_DIR, "output.csv"), index=False)
#     df.to_json(os.path.join(OUT_DIR, "output.json"), orient="records", indent=2)
#     df.to_excel(os.path.join(OUT_DIR, "output.xlsx"), index=False)
#     # omit .xls if xlwt not installed
#     try:
#         df.to_excel(os.path.join(OUT_DIR, "output.xls"), index=False, engine="xlwt")
#     except ValueError:
#         pass
#     # omit parquet if no engine
#     try:
#         df.to_parquet(os.path.join(OUT_DIR, "output.parquet"), index=False)
#     except ImportError:
#         pass

# # -----------------------------------------------------------------------------
# # 5. Emit compressed / “.pst” as just a JSON dump with .pst extension
# # -----------------------------------------------------------------------------
# def emit_compressed(records):
#     with open(os.path.join(OUT_DIR, "output.pst"), "w", encoding="utf-8") as f:
#         json.dump(records, f, ensure_ascii=False, indent=2)

# # -----------------------------------------------------------------------------
# # 6. Emit unstructured formats
# # -----------------------------------------------------------------------------
# def emit_txt(records):
#     with open(os.path.join(OUT_DIR, "output.txt"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(r["text"] + "\n\n")

# def emit_log(records):
#     with open(os.path.join(OUT_DIR, "output.log"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(f"{r['variation'].upper()}: {r['text']}\n")

# def emit_html(records):
#     with open(os.path.join(OUT_DIR, "output.html"), "w", encoding="utf-8") as f:
#         f.write("<!doctype html><html><body>\n")
#         for r in records:
#             f.write(f"<p>{r['text']}</p>\n")
#         f.write("</body></html>")

# def emit_pdf(records):
#     c = canvas.Canvas(os.path.join(OUT_DIR, "output.pdf"), pagesize=letter)
#     width, height = letter
#     y = height - 40
#     for r in records:
#         c.drawString(40, y, r["text"])
#         y -= 14
#         if y < 40:
#             c.showPage()
#             y = height - 40
#     c.save()

# def emit_docx(records):
#     doc = Document()
#     for r in records:
#         doc.add_paragraph(r["text"])
#     doc.save(os.path.join(OUT_DIR, "output.docx"))

# def emit_sql(records):
#     with open(os.path.join(OUT_DIR, "output.sql"), "w", encoding="utf-8") as f:
#         for r in records:
#             txt = r["text"].replace("'", "''")
#             var = r["variation"]
#             val = '1' if r.get("is_valid") else '0'
#             f.write(f"INSERT INTO sentences(text, variation, is_valid) VALUES('{txt}', '{var}', {val});\n")

# def emit_diff(records):
#     with open(os.path.join(OUT_DIR, "output.diff"), "w", encoding="utf-8") as f:
#         f.write("--- original\n+++ mixed\n")
#         for r in records:
#             f.write(f"+ {r['text']}\n")
#         f.write("\n")

# # -----------------------------------------------------------------------------
# # 6½. Emit real‑world paragraphs
# # -----------------------------------------------------------------------------
# def emit_paragraphs(records):
#     needed = PARA_COUNT * SENTS_PER_PARA
#     chunk  = records[:needed]
#     paras  = []
#     for i in range(PARA_COUNT):
#         start = i * SENTS_PER_PARA
#         end   = start + SENTS_PER_PARA
#         sents = [r["text"] for r in chunk[start:end]]
#         paras.append(" ".join(sents))
#     with open(os.path.join(OUT_DIR, "output_paras.txt"), "w", encoding="utf-8") as f:
#         for p in paras:
#             f.write(p + "\n\n")

# # -----------------------------------------------------------------------------
# # 7. Main entrypoint
# # -----------------------------------------------------------------------------
# if __name__ == "__main__":
#     count  = int(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_COUNT
#     records = collect_records(count)
#     df      = pd.DataFrame(records)

#     # structured
#     emit_structured(df)
#     # compressed
#     emit_compressed(records)
#     # unstructured
#     emit_txt(records)
#     emit_log(records)
#     emit_html(records)
#     emit_pdf(records)
#     emit_docx(records)
#     emit_sql(records)
#     emit_diff(records)
#     # paragraphs
#     emit_paragraphs(records)

#     print(f"✅ Generated {count} records per module from {len(MODULES)} modules")
#     print(f"   → outputs in folder: {OUT_DIR}")    




#----------------------------random but without checksum------------------------------------------
#!/usr/bin/env python3
# import importlib
# import os
# import sys
# import random
# import json
# import pandas as pd

# # PDF & DOCX
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# from docx import Document

# # -----------------------------------------------------------------------------
# # 1. Configuration
# # -----------------------------------------------------------------------------
# DEFAULT_COUNT   = 50    # sentences per module if you omit the CLI arg
# PARA_COUNT      = 5     # number of paragraphs to emit
# SENTS_PER_PARA  = 5     # sentences per paragraph

# OUT_DIR = "mixed_output"
# os.makedirs(OUT_DIR, exist_ok=True)

# # -----------------------------------------------------------------------------
# # 2. Discover & import your generate_*.py modules
# # -----------------------------------------------------------------------------
# MODULES = []
# for fn in os.listdir("."):
#     if fn.startswith("generate_") and fn.endswith(".py"):
#         MODULES.append(importlib.import_module(fn[:-3]))

# # -----------------------------------------------------------------------------
# # 3. Find each module’s *_variations() function by scanning its globals
# # -----------------------------------------------------------------------------
# def get_generator(mod):
#     funcs = [
#         getattr(mod, name)
#         for name in dir(mod)
#         if name.endswith("_variations") and callable(getattr(mod, name))
#     ]
#     if len(funcs) == 1:
#         return funcs[0]
#     elif not funcs:
#         raise AttributeError(f"No `*_variations` function found in {mod.__name__}")
#     else:
#         # if there's more than one, pick the one matching the module name
#         base = mod.__name__.split("generate_")[-1]
#         for fn in funcs:
#             if base in fn.__name__:
#                 return fn
#         # fallback
#         return funcs[0]

# # -----------------------------------------------------------------------------
# # 4. Collect a shuffled pool of records
# # -----------------------------------------------------------------------------
# def collect_records(count_per_module):
#     pool = []
#     for mod in MODULES:
#         gen = get_generator(mod)
#         pool.extend(gen(count_per_module))
#     random.shuffle(pool)
#     return pool

# # -----------------------------------------------------------------------------
# # 5. Structured output via pandas
# # -----------------------------------------------------------------------------
# def emit_structured(df):
#     df.to_csv ( os.path.join(OUT_DIR, "output.csv"), index=False )
#     df.to_json( os.path.join(OUT_DIR, "output.json"), orient="records", indent=2 )
#     df.to_excel(os.path.join(OUT_DIR, "output.xlsx"), index=False)
#     # optional XLS
#     try:
#         df.to_excel(os.path.join(OUT_DIR, "output.xls"), index=False, engine="xlwt")
#     except Exception:
#         pass
#     # optional Parquet
#     try:
#         df.to_parquet(os.path.join(OUT_DIR, "output.parquet"), index=False)
#     except Exception:
#         pass

# # -----------------------------------------------------------------------------
# # 6. Compressed dump as .pst
# # -----------------------------------------------------------------------------
# def emit_compressed(records):
#     with open(os.path.join(OUT_DIR, "output.pst"), "w", encoding="utf-8") as f:
#         json.dump(records, f, ensure_ascii=False, indent=2)

# # -----------------------------------------------------------------------------
# # 7. Unstructured formats
# # -----------------------------------------------------------------------------
# def emit_txt(records):
#     with open(os.path.join(OUT_DIR, "output.txt"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(r["text"] + "\n\n")

# def emit_log(records):
#     with open(os.path.join(OUT_DIR, "output.log"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(f"{r['variation'].upper()}: {r['text']}\n")

# def emit_html(records):
#     with open(os.path.join(OUT_DIR, "output.html"), "w", encoding="utf-8") as f:
#         f.write("<!doctype html><html><body>\n")
#         for r in records:
#             f.write(f"<p>{r['text']}</p>\n")
#         f.write("</body></html>")

# def emit_pdf(records):
#     c = canvas.Canvas(os.path.join(OUT_DIR, "output.pdf"), pagesize=letter)
#     w, h = letter
#     y = h - 40
#     for r in records:
#         c.drawString(40, y, r["text"])
#         y -= 14
#         if y < 40:
#             c.showPage()
#             y = h - 40
#     c.save()

# def emit_docx(records):
#     doc = Document()
#     for r in records:
#         doc.add_paragraph(r["text"])
#     doc.save(os.path.join(OUT_DIR, "output.docx"))

# def emit_sql(records):
#     with open(os.path.join(OUT_DIR, "output.sql"), "w", encoding="utf-8") as f:
#         for r in records:
#             txt = r["text"].replace("'", "''")
#             var = r["variation"]
#             val = '1' if r.get("is_valid") else '0'
#             f.write(f"INSERT INTO sentences(text, variation, is_valid) VALUES('{txt}', '{var}', {val});\n")

# def emit_diff(records):
#     with open(os.path.join(OUT_DIR, "output.diff"), "w", encoding="utf-8") as f:
#         f.write("--- original\n+++ mixed\n")
#         for r in records:
#             f.write(f"+ {r['text']}\n")
#         f.write("\n")

# # -----------------------------------------------------------------------------
# # 8. Paragraph emitter
# # -----------------------------------------------------------------------------
# def emit_paragraphs(records):
#     needed = PARA_COUNT * SENTS_PER_PARA
#     block  = records[:needed]
#     paras  = []
#     for i in range(PARA_COUNT):
#         chunk = block[i*SENTS_PER_PARA:(i+1)*SENTS_PER_PARA]
#         paras.append(" ".join(r["text"] for r in chunk))
#     with open(os.path.join(OUT_DIR, "output_paras.txt"), "w", encoding="utf-8") as f:
#         for p in paras:
#             f.write(p + "\n\n")


# # -----------------------------------------------------------------------------
# # 9. Main
# # -----------------------------------------------------------------------------
# if __name__ == "__main__":
#     count   = int(sys.argv[1]) if len(sys.argv)>1 else DEFAULT_COUNT
#     records = collect_records(count)
#     df      = pd.DataFrame(records)

#     emit_structured(df)
#     emit_compressed(records)
#     emit_txt(records)
#     emit_log(records)
#     emit_html(records)
#     emit_pdf(records)
#     emit_docx(records)
#     emit_sql(records)
#     emit_diff(records)
#     emit_paragraphs(records)

#     print(f"✅ Generated {count} records per module from {len(MODULES)} generators")
#     print(f"   → all outputs in folder: {OUT_DIR}")
# #----------------------------------------------------------------------------------------------


# #!/usr/bin/env python3
# import os, sys, random, importlib, json
# import pandas as pd
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# from docx import Document

# # -----------------------------------------------------------------------------
# # 1. Configuration
# # -----------------------------------------------------------------------------
# DEFAULT_COUNT   = 50    # records per module
# PARA_COUNT      = 5     # paragraphs
# SENTS_PER_PARA  = 5     # sentences each
# OUT_DIR         = "mixed_output"
# os.makedirs(OUT_DIR, exist_ok=True)

# # -----------------------------------------------------------------------------
# # 2. Discover & import generate_*.py modules
# # -----------------------------------------------------------------------------
# MODULES = []
# for fn in os.listdir("."):
#     if fn.startswith("generate_") and fn.endswith(".py"):
#         MODULES.append(importlib.import_module(fn[:-3]))

# # -----------------------------------------------------------------------------
# # 3. Locate each module’s sole *_variations() function
# # -----------------------------------------------------------------------------
# def get_generator(mod):
#     fns = [getattr(mod, n) for n in dir(mod)
#            if n.endswith("_variations") and callable(getattr(mod, n))]
#     if not fns:
#         raise AttributeError(f"No *_variations in {mod.__name__}")
#     if len(fns)==1:
#         return fns[0]
#     # disambiguate by module name
#     base = mod.__name__.split("generate_")[-1]
#     for fn in fns:
#         if base in fn.__name__:
#             return fn
#     return fns[0]

# # -----------------------------------------------------------------------------
# # 4. Build & shuffle the pool
# # -----------------------------------------------------------------------------
# def collect_records(count_per_module):
#     pool = []
#     for mod in MODULES:
#         gen = get_generator(mod)
#         pool.extend(gen(count_per_module))
#     random.shuffle(pool)
#     return pool

# # -----------------------------------------------------------------------------
# # 5. Structured output
# # -----------------------------------------------------------------------------
# def emit_structured(df):
#     df.to_csv ( os.path.join(OUT_DIR, "output.csv"), index=False )
#     df.to_json( os.path.join(OUT_DIR, "output.json"), orient="records", indent=2 )
#     df.to_excel(os.path.join(OUT_DIR, "output.xlsx"), index=False)
#     try:
#         df.to_excel(os.path.join(OUT_DIR, "output.xls"), index=False, engine="xlwt")
#     except ImportError:
#         pass
#     try:
#         df.to_parquet(os.path.join(OUT_DIR, "output.parquet"), index=False)
#     except ImportError:
#         pass

# # -----------------------------------------------------------------------------
# # 6. Compressed (.pst)
# # -----------------------------------------------------------------------------
# def emit_compressed(records):
#     with open(os.path.join(OUT_DIR, "output.pst"), "w", encoding="utf-8") as f:
#         json.dump(records, f, ensure_ascii=False, indent=2)

# # -----------------------------------------------------------------------------
# # 7. Unstructured
# # -----------------------------------------------------------------------------
# def emit_txt(records):
#     with open(os.path.join(OUT_DIR, "output.txt"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(r["text"] + "\n\n")

# def emit_log(records):
#     with open(os.path.join(OUT_DIR, "output.log"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(f"{r['variation'].upper()}: {r['text']}\n")

# def emit_html(records):
#     with open(os.path.join(OUT_DIR, "output.html"), "w", encoding="utf-8") as f:
#         f.write("<!doctype html><html><body>\n")
#         for r in records:
#             f.write(f"<p>{r['text']}</p>\n")
#         f.write("</body></html>")

# def emit_pdf(records):
#     c = canvas.Canvas(os.path.join(OUT_DIR, "output.pdf"), pagesize=letter)
#     w,h = letter; y = h-40
#     for r in records:
#         c.drawString(40, y, r["text"])
#         y -= 14
#         if y < 40:
#             c.showPage(); y = h-40
#     c.save()

# def emit_docx(records):
#     doc = Document()
#     for r in records:
#         doc.add_paragraph(r["text"])
#     doc.save(os.path.join(OUT_DIR, "output.docx"))

# def emit_sql(records):
#     with open(os.path.join(OUT_DIR, "output.sql"), "w", encoding="utf-8") as f:
#         for r in records:
#             txt = r["text"].replace("'", "''")
#             var = r["variation"]
#             val = '1' if r.get("is_valid") else '0'
#             f.write(f"INSERT INTO sentences(text,variation,is_valid) VALUES('{txt}','{var}',{val});\n")

# def emit_diff(records):
#     with open(os.path.join(OUT_DIR, "output.diff"), "w", encoding="utf-8") as f:
#         f.write("--- original\n+++ mixed\n")
#         for r in records:
#             f.write(f"+ {r['text']}\n")
#         f.write("\n")

# # -----------------------------------------------------------------------------
# # 8. Paragraphs (.txt)
# # -----------------------------------------------------------------------------
# def emit_paragraphs(records):
#     needed = PARA_COUNT * SENTS_PER_PARA
#     block  = records[:needed]
#     paras  = []
#     for i in range(PARA_COUNT):
#         chunk = block[i*SENTS_PER_PARA:(i+1)*SENTS_PER_PARA]
#         paras.append(" ".join(r["text"] for r in chunk))
#     with open(os.path.join(OUT_DIR, "output_paras.txt"), "w", encoding="utf-8") as f:
#         for p in paras:
#             f.write(p + "\n\n")

# # -----------------------------------------------------------------------------
# # 9. Main
# # -----------------------------------------------------------------------------
# if __name__ == "__main__":
#     count   = int(sys.argv[1]) if len(sys.argv)>1 else DEFAULT_COUNT
#     records = collect_records(count)
#     df      = pd.DataFrame(records)

#     emit_structured(df)
#     emit_compressed(records)
#     emit_txt(records)
#     emit_log(records)
#     emit_html(records)
#     emit_pdf(records)
#     emit_docx(records)
#     emit_sql(records)
#     emit_diff(records)
#     emit_paragraphs(records)

#     print(f"✅ Generated {count} records per module from {len(MODULES)} generators")
#     print(f"   → all outputs in folder: {OUT_DIR}")







# #!/usr/bin/env python3
# import importlib
# import os
# import sys
# import random
# import json
# import pandas as pd

# # PDF & DOCX
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# from docx import Document

# # -----------------------------------------------------------------------------
# # 1. Configuration
# # -----------------------------------------------------------------------------
# DEFAULT_COUNT   = 50    # sentences per module if you omit the CLI arg
# PARA_COUNT      = 5     # number of paragraphs to emit
# SENTS_PER_PARA  = 5     # sentences per paragraph
# OUT_DIR         = "mixed_output"
# os.makedirs(OUT_DIR, exist_ok=True)

# # -----------------------------------------------------------------------------
# # 2. Discover & import your generate_*.py modules
# # -----------------------------------------------------------------------------
# MODULES = []
# for fn in os.listdir("."):
#     if fn.startswith("generate_") and fn.endswith(".py"):
#         MODULES.append(importlib.import_module(fn[:-3]))

# # -----------------------------------------------------------------------------
# # 3. Find each module’s *_variations() function by scanning its globals
# # -----------------------------------------------------------------------------
# def get_generator(mod):
#     funcs = [
#         getattr(mod, name)
#         for name in dir(mod)
#         if name.endswith("_variations") and callable(getattr(mod, name))
#     ]
#     if len(funcs) == 1:
#         return funcs[0]
#     if not funcs:
#         raise AttributeError(f"No `*_variations` function found in {mod.__name__}")
#     # disambiguate by module name
#     base = mod.__name__.split("generate_")[-1]
#     for fn in funcs:
#         if base in fn.__name__:
#             return fn
#     return funcs[0]

# # -----------------------------------------------------------------------------
# # 4. Collect a shuffled pool of records
# # -----------------------------------------------------------------------------
# def collect_records(count_per_module):
#     pool = []
#     for mod in MODULES:
#         gen = get_generator(mod)
#         pool.extend(gen(count_per_module))
#     random.shuffle(pool)
#     return pool

# # -----------------------------------------------------------------------------
# # 5. Structured output via pandas
# # -----------------------------------------------------------------------------
# def emit_structured(df):
#     # CSV
#     df.to_csv(os.path.join(OUT_DIR, "output.csv"), index=False)
#     # JSON
#     df.to_json(os.path.join(OUT_DIR, "output.json"), orient="records", indent=2)
#     # Real XLSX
#     df.to_excel(os.path.join(OUT_DIR, "output.xlsx"), index=False)
#     # “Fake” XLS: write as XLSX but with .xls extension via openpyxl
#     df.to_excel(
#         os.path.join(OUT_DIR, "output.xls"),
#         index=False,
#         engine="openpyxl"
#     )
#     # Parquet if available
#     try:
#         df.to_parquet(os.path.join(OUT_DIR, "output.parquet"), index=False)
#     except ImportError:
#         pass

# # -----------------------------------------------------------------------------
# # 6. Compressed dump as .pst
# # -----------------------------------------------------------------------------
# def emit_compressed(records):
#     with open(os.path.join(OUT_DIR, "output.pst"), "w", encoding="utf-8") as f:
#         json.dump(records, f, ensure_ascii=False, indent=2)

# # -----------------------------------------------------------------------------
# # 7. Unstructured formats
# # -----------------------------------------------------------------------------
# def emit_txt(records):
#     with open(os.path.join(OUT_DIR, "output.txt"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(r["text"] + "\n\n")

# def emit_log(records):
#     with open(os.path.join(OUT_DIR, "output.log"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(f"{r['variation'].upper()}: {r['text']}\n")

# def emit_html(records):
#     with open(os.path.join(OUT_DIR, "output.html"), "w", encoding="utf-8") as f:
#         f.write("<!doctype html><html><body>\n")
#         for r in records:
#             f.write(f"<p>{r['text']}</p>\n")
#         f.write("</body></html>")

# def emit_pdf(records):
#     c = canvas.Canvas(os.path.join(OUT_DIR, "output.pdf"), pagesize=letter)
#     w, h = letter
#     y = h - 40
#     for r in records:
#         c.drawString(40, y, r["text"])
#         y -= 14
#         if y < 40:
#             c.showPage()
#             y = h - 40
#     c.save()

# def emit_docx(records):
#     doc = Document()
#     for r in records:
#         doc.add_paragraph(r["text"])
#     doc.save(os.path.join(OUT_DIR, "output.docx"))

# def emit_sql(records):
#     with open(os.path.join(OUT_DIR, "output.sql"), "w", encoding="utf-8") as f:
#         for r in records:
#             txt = r["text"].replace("'", "''")
#             var = r["variation"]
#             val = '1' if r.get("is_valid") else '0'
#             f.write(
#                 "INSERT INTO sentences(text, variation, is_valid) "
#                 f"VALUES('{txt}', '{var}', {val});\n"
#             )

# def emit_diff(records):
#     with open(os.path.join(OUT_DIR, "output.diff"), "w", encoding="utf-8") as f:
#         f.write("--- original\n+++ mixed\n")
#         for r in records:
#             f.write(f"+ {r['text']}\n")
#         f.write("\n")

# # -----------------------------------------------------------------------------
# # 8. Paragraph emitter
# # -----------------------------------------------------------------------------
# def emit_paragraphs(records):
#     needed = PARA_COUNT * SENTS_PER_PARA
#     block  = records[:needed]
#     paras  = []
#     for i in range(PARA_COUNT):
#         chunk = block[i*SENTS_PER_PARA:(i+1)*SENTS_PER_PARA]
#         paras.append(" ".join(r["text"] for r in chunk))
#     with open(os.path.join(OUT_DIR, "output_paras.txt"), "w", encoding="utf-8") as f:
#         for p in paras:
#             f.write(p + "\n\n")

# # -----------------------------------------------------------------------------
# # 9. Main
# # -----------------------------------------------------------------------------
# if __name__ == "__main__":
#     count   = int(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_COUNT
#     records = collect_records(count)
#     df      = pd.DataFrame(records)

#     emit_structured(df)
#     emit_compressed(records)
#     emit_txt(records)
#     emit_log(records)
#     emit_html(records)
#     emit_pdf(records)
#     emit_docx(records)
#     emit_sql(records)
#     emit_diff(records)
#     emit_paragraphs(records)

#     print(f"✅ Generated {count} records per module from {len(MODULES)} generators")
#     print(f"   → all outputs in folder: {OUT_DIR}")





# #!/usr/bin/env python3
# import importlib
# import os
# import sys
# import random
# import json
# import gzip
# import pandas as pd

# # PDF & DOCX
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# from docx import Document

# # -----------------------------------------------------------------------------
# # 1. Configuration
# # -----------------------------------------------------------------------------
# DEFAULT_COUNT   = 50    # sentences per module if you omit the CLI arg
# PARA_COUNT      = 5     # number of paragraphs to emit
# SENTS_PER_PARA  = 5     # sentences per paragraph
# OUT_DIR         = "mixed_output"
# os.makedirs(OUT_DIR, exist_ok=True)

# # -----------------------------------------------------------------------------
# # 2. Discover & import your generate_*.py modules
# # -----------------------------------------------------------------------------
# MODULES = []
# for fn in os.listdir("."):
#     if fn.startswith("generate_") and fn.endswith(".py"):
#         MODULES.append(importlib.import_module(fn[:-3]))

# # -----------------------------------------------------------------------------
# # 3. Find each module’s *_variations() function
# # -----------------------------------------------------------------------------
# def get_generator(mod):
#     funcs = [
#         getattr(mod, name)
#         for name in dir(mod)
#         if name.endswith("_variations") and callable(getattr(mod, name))
#     ]
#     if len(funcs) == 1:
#         return funcs[0]
#     if not funcs:
#         raise AttributeError(f"No `*_variations` in {mod.__name__}")
#     base = mod.__name__.split("generate_")[-1]
#     for fn in funcs:
#         if base in fn.__name__:
#             return fn
#     return funcs[0]

# # -----------------------------------------------------------------------------
# # 4. Build & shuffle the pool
# # -----------------------------------------------------------------------------
# def collect_records(count_per_module):
#     pool = []
#     for mod in MODULES:
#         gen = get_generator(mod)
#         pool.extend(gen(count_per_module))
#     random.shuffle(pool)
#     return pool

# # -----------------------------------------------------------------------------
# # 5. Structured output
# # -----------------------------------------------------------------------------
# def emit_structured(df, records):
#     # 5.1 CSV
#     df.to_csv(os.path.join(OUT_DIR, "output.csv"), index=False)

#     # 5.2 JSON without nulls
#     cleaned = [
#         {k: v for k, v in r.items() if v is not None}
#         for r in records
#     ]
#     with open(os.path.join(OUT_DIR, "output.json"), "w", encoding="utf-8") as f:
#         json.dump(cleaned, f, ensure_ascii=False, indent=2)

#     # 5.3 XLSX
#     df.to_excel(os.path.join(OUT_DIR, "output.xlsx"), index=False)

#     # 5.4 “.xls” via openpyxl so it never errors
#     df.to_excel(
#         os.path.join(OUT_DIR, "output.xls"),
#         index=False,
#         engine="openpyxl"
#     )

#     # 5.5 Parquet if available
#     try:
#         df.to_parquet(os.path.join(OUT_DIR, "output.parquet"), index=False)
#     except ImportError:
#         pass

# # -----------------------------------------------------------------------------
# # 6. Compressed dump as true gzip‑.pst
# # -----------------------------------------------------------------------------
# def emit_compressed(records):
#     cleaned = [
#         {k: v for k, v in r.items() if v is not None}
#         for r in records
#     ]
#     gz_path = os.path.join(OUT_DIR, "output.pst")
#     with gzip.open(gz_path, "wt", encoding="utf-8") as f:
#         json.dump(cleaned, f, ensure_ascii=False)

# # -----------------------------------------------------------------------------
# # 7. Unstructured formats
# # -----------------------------------------------------------------------------
# def emit_txt(records):
#     with open(os.path.join(OUT_DIR, "output.txt"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(r["text"] + "\n\n")

# def emit_log(records):
#     with open(os.path.join(OUT_DIR, "output.log"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(f"{r['variation'].upper()}: {r['text']}\n")

# def emit_html(records):
#     with open(os.path.join(OUT_DIR, "output.html"), "w", encoding="utf-8") as f:
#         f.write("<!doctype html><html><body>\n")
#         for r in records:
#             f.write(f"<p>{r['text']}</p>\n")
#         f.write("</body></html>")

# def emit_pdf(records):
#     c = canvas.Canvas(os.path.join(OUT_DIR, "output.pdf"), pagesize=letter)
#     w, h = letter
#     y = h - 40
#     for r in records:
#         c.drawString(40, y, r["text"])
#         y -= 14
#         if y < 40:
#             c.showPage()
#             y = h - 40
#     c.save()

# def emit_docx(records):
#     doc = Document()
#     for r in records:
#         doc.add_paragraph(r["text"])
#     doc.save(os.path.join(OUT_DIR, "output.docx"))

# def emit_sql(records):
#     with open(os.path.join(OUT_DIR, "output.sql"), "w", encoding="utf-8") as f:
#         for r in records:
#             txt = r["text"].replace("'", "''")
#             var = r["variation"]
#             val = "1" if r.get("is_valid") else "0"
#             f.write(
#                 "INSERT INTO sentences(text,variation,is_valid) "
#                 f"VALUES('{txt}','{var}',{val});\n"
#             )

# def emit_diff(records):
#     with open(os.path.join(OUT_DIR, "output.diff"), "w", encoding="utf-8") as f:
#         f.write("--- original\n+++ mixed\n")
#         for r in records:
#             f.write(f"+ {r['text']}\n")
#         f.write("\n")

# # -----------------------------------------------------------------------------
# # 8. Paragraph emitter
# # -----------------------------------------------------------------------------
# def emit_paragraphs(records):
#     needed = PARA_COUNT * SENTS_PER_PARA
#     block  = records[:needed]
#     paras  = []
#     for i in range(PARA_COUNT):
#         chunk = block[i*SENTS_PER_PARA:(i+1)*SENTS_PER_PARA]
#         paras.append(" ".join(r["text"] for r in chunk))
#     with open(os.path.join(OUT_DIR, "output_paras.txt"), "w", encoding="utf-8") as f:
#         for p in paras:
#             f.write(p + "\n\n")

# # -----------------------------------------------------------------------------
# # 9. Main
# # -----------------------------------------------------------------------------
# if __name__ == "__main__":
#     count   = int(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_COUNT
#     records = collect_records(count)
#     df      = pd.DataFrame(records)

#     emit_structured(df, records)
#     emit_compressed(records)
#     emit_txt(records)
#     emit_log(records)
#     emit_html(records)
#     emit_pdf(records)
#     emit_docx(records)
#     emit_sql(records)
#     emit_diff(records)
#     emit_paragraphs(records)

#     print(f"✅ Generated {count} records per module from {len(MODULES)} generators")
#     print(f"   → all outputs in folder: {OUT_DIR}")



# #!/usr/bin/env python3
# import os
# import sys
# import random
# import json
# import re
# import importlib
# from datetime import datetime
# import pandas as pd

# # stdnum imports
# from stdnum import luhn
# from stdnum.in_ import pan, gstin, epic

# # PDF & DOCX
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# from docx import Document

# # -----------------------------------------------------------------------------
# # 1. Configuration
# # -----------------------------------------------------------------------------
# DEFAULT_COUNT   = 50    # records per module if you omit the CLI arg
# PARA_COUNT      = 5     # paragraphs to emit
# SENTS_PER_PARA  = 5     # sentences per paragraph
# OUT_DIR         = "mixed_output"
# os.makedirs(OUT_DIR, exist_ok=True)

# # -----------------------------------------------------------------------------
# # 2. Discover & import your generate_*.py modules
# # -----------------------------------------------------------------------------
# MODULES = []
# for fn in os.listdir("."):
#     if fn.startswith("generate_") and fn.endswith(".py"):
#         MODULES.append(importlib.import_module(fn[:-3]))

# # -----------------------------------------------------------------------------
# # 3. Find each module’s *_variations() generator
# # -----------------------------------------------------------------------------
# def get_generator(mod):
#     fns = [getattr(mod,n) for n in dir(mod)
#            if n.endswith("_variations") and callable(getattr(mod,n))]
#     if not fns:
#         raise AttributeError(f"No *_variations in {mod.__name__}")
#     if len(fns)==1:
#         return fns[0]
#     base = mod.__name__.split("generate_")[-1]
#     for fn in fns:
#         if base in fn.__name__:
#             return fn
#     return fns[0]

# # -----------------------------------------------------------------------------
# # 4. Build & shuffle the pool
# # -----------------------------------------------------------------------------
# def collect_records(count_per_module):
#     pool = []
#     for mod in MODULES:
#         gen = get_generator(mod)
#         pool.extend(gen(count_per_module))
#     random.shuffle(pool)
#     return pool

# # -----------------------------------------------------------------------------
# # 5. Validators (always strip/normalize then validate)
# # -----------------------------------------------------------------------------
# def validate_account(acc):
#     digits = re.sub(r"\D", "", acc)
#     return len(digits)==14

# def validate_card(card):
#     digits = re.sub(r"\D", "", card)
#     return luhn.is_valid(digits)

# def validate_dob(s):
#     for fmt in ("%d/%m/%Y","%d-%m-%Y","%Y-%m-%d","%d %b %Y"):
#         try:
#             datetime.strptime(s, fmt)
#             return True
#         except ValueError:
#             continue
#     return False

# def validate_ifsc(code):
#     c = re.sub(r"[\s-]", "", code).upper()
#     return bool(re.fullmatch(r"^[A-Z]{4}0\d{6}$", c))

# def validate_imei(imei):
#     digits = re.sub(r"\D", "", imei)
#     if len(digits)!=15: return False
#     return luhn.is_valid(digits)

# def validate_ip(ip):
#     return bool(re.fullmatch(r"(25[0-5]|2[0-4]\d|1\d\d|\d\d?)\."
#                              r"(25[0-5]|2[0-4]\d|1\d\d|\d\d?)\."
#                              r"(25[0-5]|2[0-4]\d|1\d\d|\d\d?)\."
#                              r"(25[0-5]|2[0-4]\d|1\d\d|\d\d?)", ip))

# def validate_mac(mac):
#     m = mac.replace("-",":").upper()
#     return bool(re.fullmatch(r"([0-9A-F]{2}:){5}[0-9A-F]{2}", m))

# def validate_email(e):
#     return bool(re.fullmatch(r"^[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}$", e))

# def validate_gender(g):
#     return g.lower() in ("male","female","other")

# def validate_gstin(g):
#     return gstin.is_valid(re.sub(r"\W","",g).upper())

# def validate_pan_fn(p):
#     return pan.is_valid(p)

# def validate_passport(p):
#     return bool(re.fullmatch(r"^[A-Z][0-9]{7}$", p))

# def validate_phone(p):
#     d = re.sub(r"\D", "", p)
#     if len(d)==10 and d[0] in "6789": return True
#     if len(d)==12 and d.startswith("91") and d[2] in "6789": return True
#     return False

# def validate_ration(r):
#     # very loose: 5–16 alnum, spaces/dashes allowed
#     return bool(re.fullmatch(r"[A-Za-z0-9 \-]{5,16}", r))

# def validate_upi(u):
#     return bool(re.fullmatch(r"^[A-Za-z0-9._]{3,256}@[A-Za-z]{3,64}$", u))

# def validate_vehicle(v):
#     return bool(re.fullmatch(r"^[A-Z]{2}\d{2}[A-Z]{2}\d{4}$", v))

# def validate_epic(e):
#     return epic.is_valid(re.sub(r"\W","", e).upper())

# # -----------------------------------------------------------------------------
# # 6. Writers
# # -----------------------------------------------------------------------------
# def emit_structured(df):
#     df.to_csv ( os.path.join(OUT_DIR, "output.csv"), index=False )
#     df.to_json( os.path.join(OUT_DIR, "output.json"), orient="records", indent=2 )
#     df.to_excel(os.path.join(OUT_DIR, "output.xlsx"), index=False)
#     try:
#         df.to_excel(os.path.join(OUT_DIR, "output.xls"), index=False, engine="xlwt")
#     except Exception:
#         pass
#     try:
#         df.to_parquet(os.path.join(OUT_DIR, "output.parquet"), index=False)
#     except Exception:
#         pass

# def emit_compressed(records):
#     with open(os.path.join(OUT_DIR, "output.pst"), "w", encoding="utf-8") as f:
#         json.dump(records, f, ensure_ascii=False, indent=2)

# def emit_txt(records):
#     with open(os.path.join(OUT_DIR, "output.txt"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(r["text"] + "\n\n")

# def emit_log(records):
#     with open(os.path.join(OUT_DIR, "output.log"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(f"{r['variation'].upper()}: {r['text']}\n")

# def emit_html(records):
#     # embed valid JSON in a <pre> for browser viewing
#     j = json.dumps(records, indent=2)
#     with open(os.path.join(OUT_DIR, "output.html"), "w", encoding="utf-8") as f:
#         f.write("<!doctype html><html><body><pre>\n")
#         f.write(j)
#         f.write("\n</pre></body></html>")

# def emit_pdf(records):
#     c = canvas.Canvas(os.path.join(OUT_DIR, "output.pdf"), pagesize=letter)
#     w,h = letter; y = h - 40
#     for r in records:
#         c.drawString(40, y, r["text"])
#         y -= 14
#         if y < 40:
#             c.showPage(); y = h - 40
#     c.save()

# def emit_rtf(records):
#     # write a simple RTF (.doc) that Word will open
#     with open(os.path.join(OUT_DIR, "output.doc"), "w", encoding="utf-8") as f:
#         f.write(r"{\rtf1\ansi"+"\n")
#         for r in records:
#             txt = r["text"].replace("\\","\\\\").replace("{","\\{").replace("}","\\}")
#             f.write(txt + r"\par" + "\n")
#         f.write("}")

# def emit_docx(records):
#     doc = Document()
#     for r in records:
#         doc.add_paragraph(json.dumps(r, ensure_ascii=False))
#     doc.save(os.path.join(OUT_DIR, "output.docx"))

# def emit_sql(records):
#     with open(os.path.join(OUT_DIR, "output.sql"), "w", encoding="utf-8") as f:
#         for r in records:
#             txt = r["text"].replace("'", "''")
#             var = r["variation"]
#             val = '1' if r.get("is_valid") else '0'
#             f.write(
#               f"INSERT INTO sentences(text,variation,is_valid) "
#               f"VALUES('{txt}','{var}',{val});\n"
#             )

# def emit_diff(records):
#     with open(os.path.join(OUT_DIR, "output.diff"), "w", encoding="utf-8") as f:
#         f.write("--- original\n+++ mixed\n")
#         for r in records:
#             f.write(f"+ {r['text']}\n")

# def emit_paragraphs(records):
#     needed = PARA_COUNT * SENTS_PER_PARA
#     block  = records[:needed]
#     paras  = []
#     for i in range(PARA_COUNT):
#         chunk = block[i*SENTS_PER_PARA:(i+1)*SENTS_PER_PARA]
#         paras.append(" ".join(r["text"] for r in chunk))
#     with open(os.path.join(OUT_DIR, "output_paras.txt"), "w", encoding="utf-8") as f:
#         for p in paras:
#             f.write(p + "\n\n")

# # -----------------------------------------------------------------------------
# # 7. Main
# # -----------------------------------------------------------------------------
# if __name__ == "__main__":
#     count   = int(sys.argv[1]) if len(sys.argv)>1 else DEFAULT_COUNT
#     records = collect_records(count)

#     # re-validate each PII field (in-place) so that `is_valid` flags
#     for r in records:
#         # pick whichever field is present and re-check
#         if r.get("card")        : r["is_valid"] = validate_card(r["card"])
#         if r.get("account")     : r["is_valid"] = validate_account(r["account"])
#         if r.get("dob")         : r["is_valid"] = validate_dob(r["dob"])
#         if r.get("ifsc")        : r["is_valid"] = validate_ifsc(r["ifsc"])
#         if r.get("imei")        : r["is_valid"] = validate_imei(r["imei"])
#         if r.get("ip")          : r["is_valid"] = validate_ip(r["ip"])
#         if r.get("mac")         : r["is_valid"] = validate_mac(r["mac"])
#         if r.get("email")       : r["is_valid"] = validate_email(r["email"])
#         if r.get("gender")      : r["is_valid"] = validate_gender(r["gender"])
#         if r.get("gstin")       : r["is_valid"] = validate_gstin(r["gstin"])
#         if r.get("pan")         : r["is_valid"] = validate_pan_fn(r["pan"])
#         if r.get("passport")    : r["is_valid"] = validate_passport(r["passport"])
#         if r.get("phone")       : r["is_valid"] = validate_phone(r["phone"])
#         if r.get("ration_card") : r["is_valid"] = validate_ration(r["ration_card"])
#         if r.get("upi_id")      : r["is_valid"] = validate_upi(r["upi_id"])
#         if r.get("voter_id")    : r["is_valid"] = validate_epic(r["voter_id"])
#         if r.get("registration"): r["is_valid"] = validate_vehicle(r["registration"])

#     df = pd.DataFrame(records)

#     emit_structured(df)
#     emit_compressed(records)
#     emit_txt(records)
#     emit_log(records)
#     emit_html(records)
#     emit_pdf(records)
#     emit_rtf(records)     # produces output.doc (RTF)
#     emit_docx(records)
#     emit_sql(records)
#     emit_diff(records)
#     emit_paragraphs(records)

#     print(f"✅ Generated {count} records per module from {len(MODULES)} generators")
#     print(f"   → all outputs in folder: {OUT_DIR}")




# #!/usr/bin/env python3
# import importlib
# import os
# import sys
# import random
# import json
# import pandas as pd

# # PDF & DOCX
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# from docx import Document

# # -----------------------------------------------------------------------------
# # 1. Configuration
# # -----------------------------------------------------------------------------
# DEFAULT_COUNT   = 50    # records per module if you omit the CLI arg
# PARA_COUNT      = 5     # number of paragraphs to emit
# SENTS_PER_PARA  = 5     # sentences per paragraph

# OUT_DIR = "mixed_output"
# os.makedirs(OUT_DIR, exist_ok=True)

# # -----------------------------------------------------------------------------
# # 2. Discover & import your generate_*.py modules
# # -----------------------------------------------------------------------------
# MODULES = []
# for fn in os.listdir("."):
#     if fn.startswith("generate_") and fn.endswith(".py"):
#         MODULES.append(importlib.import_module(fn[:-3]))

# # -----------------------------------------------------------------------------
# # 3. Find each module’s *_variations() function
# # -----------------------------------------------------------------------------
# def get_generator(mod):
#     funcs = [
#         getattr(mod, name)
#         for name in dir(mod)
#         if name.endswith("_variations") and callable(getattr(mod, name))
#     ]
#     if len(funcs) == 1:
#         return funcs[0]
#     if not funcs:
#         raise AttributeError(f"No `*_variations` in {mod.__name__}")
#     base = mod.__name__.split("generate_")[-1]
#     for fn in funcs:
#         if base in fn.__name__:
#             return fn
#     return funcs[0]

# # -----------------------------------------------------------------------------
# # 4. Collect a shuffled pool of records
# # -----------------------------------------------------------------------------
# def collect_records(count_per_module):
#     pool = []
#     for mod in MODULES:
#         gen = get_generator(mod)
#         pool.extend(gen(count_per_module))
#     random.shuffle(pool)
#     return pool

# # -----------------------------------------------------------------------------
# # 5. Emit all structured files + clean JSON/PST
# # -----------------------------------------------------------------------------
# def emit_structured(df, clean_recs):
#     # CSV
#     df.to_csv(os.path.join(OUT_DIR, "output.csv"), index=False)
#     # JSON (no nulls)
#     with open(os.path.join(OUT_DIR, "output.json"), "w", encoding="utf-8") as f:
#         json.dump(clean_recs, f, ensure_ascii=False, indent=2)
#     # XLSX
#     df.to_excel(os.path.join(OUT_DIR, "output.xlsx"), index=False)
#     # XLS
#     try:
#         df.to_excel(os.path.join(OUT_DIR, "output.xls"), index=False, engine="xlwt")
#     except Exception:
#         pass
#     # Parquet
#     try:
#         df.to_parquet(os.path.join(OUT_DIR, "output.parquet"), index=False)
#     except Exception:
#         pass

# def emit_compressed(clean_recs):
#     # PST as JSON
#     with open(os.path.join(OUT_DIR, "output.pst"), "w", encoding="utf-8") as f:
#         json.dump(clean_recs, f, ensure_ascii=False, indent=2)

# # -----------------------------------------------------------------------------
# # 6. Unstructured formats
# # -----------------------------------------------------------------------------
# def emit_txt(records):
#     with open(os.path.join(OUT_DIR, "output.txt"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(r["text"] + "\n\n")

# def emit_log(records):
#     with open(os.path.join(OUT_DIR, "output.log"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(f"{r['variation'].upper()}: {r['text']}\n")

# def emit_html(records):
#     with open(os.path.join(OUT_DIR, "output.html"), "w", encoding="utf-8") as f:
#         f.write("<!doctype html><html><body>\n")
#         for r in records:
#             f.write(f"<p>{r['text']}</p>\n")
#         f.write("</body></html>")

# def emit_pdf(records):
#     c = canvas.Canvas(os.path.join(OUT_DIR, "output.pdf"), pagesize=letter)
#     w, h = letter
#     y = h - 40
#     for r in records:
#         c.drawString(40, y, r["text"])
#         y -= 14
#         if y < 40:
#             c.showPage()
#             y = h - 40
#     c.save()

# def emit_doc(records):
#     # plain .doc via .docx library
#     doc = Document()
#     for r in records:
#         doc.add_paragraph(r["text"])
#     doc.save(os.path.join(OUT_DIR, "output.doc"))

# def emit_docx(records):
#     doc = Document()
#     for r in records:
#         doc.add_paragraph(r["text"])
#     doc.save(os.path.join(OUT_DIR, "output.docx"))

# def emit_sql(records):
#     with open(os.path.join(OUT_DIR, "output.sql"), "w", encoding="utf-8") as f:
#         for r in records:
#             txt = r["text"].replace("'", "''")
#             var = r["variation"]
#             isv = '1' if r.get("is_valid") else '0'
#             f.write(
#                 "INSERT INTO sentences(text, variation, is_valid) "
#                 f"VALUES('{txt}', '{var}', {isv});\n"
#             )

# def emit_diff(records):
#     with open(os.path.join(OUT_DIR, "output.diff"), "w", encoding="utf-8") as f:
#         f.write("--- original\n+++ mixed\n")
#         for r in records:
#             f.write(f"+ {r['text']}\n")
#         f.write("\n")

# # -----------------------------------------------------------------------------
# # 7. Paragraph emitter
# # -----------------------------------------------------------------------------
# def emit_paragraphs(records):
#     needed = PARA_COUNT * SENTS_PER_PARA
#     block  = records[:needed]
#     paras  = []
#     for i in range(PARA_COUNT):
#         chunk = block[i*SENTS_PER_PARA:(i+1)*SENTS_PER_PARA]
#         paras.append(" ".join(r["text"] for r in chunk))
#     with open(os.path.join(OUT_DIR, "output_paras.txt"), "w", encoding="utf-8") as f:
#         for p in paras:
#             f.write(p + "\n\n")

# # -----------------------------------------------------------------------------
# # 8. Main
# # -----------------------------------------------------------------------------
# if __name__ == "__main__":
#     count = int(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_COUNT

#     # 1) Collect & shuffle
#     records = collect_records(count)

#     # 2) Build clean JSON/PST list (drop null fields)
#     clean_recs = []
#     for r in records:
#         cr = {k: v for k, v in r.items() if v is not None}
#         clean_recs.append(cr)

#     # 3) DataFrame for structured
#     df = pd.DataFrame(records)

#     # 4) Emit all
#     emit_structured(df, clean_recs)
#     emit_compressed(clean_recs)

#     emit_txt(records)
#     emit_log(records)
#     emit_html(records)
#     emit_pdf(records)
#     emit_doc(records)
#     emit_docx(records)
#     emit_sql(records)
#     emit_diff(records)
#     emit_paragraphs(records)

#     print(f"✅ Generated {count} records per module from {len(MODULES)} generators")
#     print(f"   → all outputs in folder: {OUT_DIR}")



#!/usr/bin/env python3
# import os
# import sys
# import random
# import importlib
# import json
# import pandas as pd

# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# from docx import Document

# # -----------------------------------------------------------------------------
# # 1. Configuration
# # -----------------------------------------------------------------------------
# DEFAULT_COUNT   = 50    # records per module
# PARA_COUNT      = 5     # paragraphs
# SENTS_PER_PARA  = 5     # sentences each
# OUT_DIR         = "mixed_output"
# os.makedirs(OUT_DIR, exist_ok=True)

# # -----------------------------------------------------------------------------
# # 2. Discover & import generate_*.py modules
# # -----------------------------------------------------------------------------
# MODULES = []
# for fn in os.listdir("."):
#     if fn.startswith("generate_") and fn.endswith(".py"):
#         MODULES.append(importlib.import_module(fn[:-3]))

# # -----------------------------------------------------------------------------
# # 3. Locate each module’s *_variations() function
# # -----------------------------------------------------------------------------
# def get_generator(mod):
#     fns = [getattr(mod, n) for n in dir(mod)
#            if n.endswith("_variations") and callable(getattr(mod, n))]
#     if not fns:
#         raise AttributeError(f"No *_variations in {mod.__name__}")
#     if len(fns)==1:
#         return fns[0]
#     base = mod.__name__.split("generate_")[-1]
#     for fn in fns:
#         if base in fn.__name__:
#             return fn
#     return fns[0]

# # -----------------------------------------------------------------------------
# # 4. Build & shuffle the pool
# # -----------------------------------------------------------------------------
# def collect_records(count_per_module):
#     pool = []
#     for mod in MODULES:
#         gen = get_generator(mod)
#         pool.extend(gen(count_per_module))
#     random.shuffle(pool)
#     return pool

# # -----------------------------------------------------------------------------
# # 5. Structured output
# # -----------------------------------------------------------------------------
# def emit_structured(df):
#     df.to_csv   ( os.path.join(OUT_DIR, "output.csv"), index=False )
#     df.to_json  ( os.path.join(OUT_DIR, "output.json"), orient="records", indent=2 )
#     df.to_excel ( os.path.join(OUT_DIR, "output.xlsx"), index=False )
#     try:
#         df.to_excel(os.path.join(OUT_DIR, "output.xls"), index=False, engine="xlwt")
#     except ImportError:
#         pass
#     try:
#         df.to_parquet(os.path.join(OUT_DIR, "output.parquet"), index=False)
#     except ImportError:
#         pass

# # -----------------------------------------------------------------------------
# # 6. “PST” = JSON dump
# # -----------------------------------------------------------------------------
# def emit_pst(records):
#     # NOTE: This is *not* a real Outlook PST, but a JSON dump with .pst extension
#     with open(os.path.join(OUT_DIR, "output.pst"), "w", encoding="utf-8") as f:
#         json.dump(records, f, ensure_ascii=False, indent=2)

# # -----------------------------------------------------------------------------
# # 7. Unstructured formats
# # -----------------------------------------------------------------------------
# def emit_txt(records):
#     with open(os.path.join(OUT_DIR, "output.txt"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(r["text"] + "\n\n")

# def emit_log(records):
#     with open(os.path.join(OUT_DIR, "output.log"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(f"{r['variation'].upper()}: {r['text']}\n")

# def emit_html(records):
#     with open(os.path.join(OUT_DIR, "output.html"), "w", encoding="utf-8") as f:
#         f.write("<!doctype html>\n<html>\n  <head><meta charset='utf-8'><title>Mixed Output</title></head>\n  <body>\n")
#         for r in records:
#             # escape HTML special chars if needed
#             text = (r["text"]
#                     .replace("&", "&amp;")
#                     .replace("<", "&lt;")
#                     .replace(">", "&gt;"))
#             f.write(f"    <p>{text}</p>\n")
#         f.write("  </body>\n</html>")

# def emit_pdf(records):
#     c = canvas.Canvas(os.path.join(OUT_DIR, "output.pdf"), pagesize=letter)
#     w, h = letter; y = h - 40
#     for r in records:
#         c.drawString(40, y, r["text"])
#         y -= 14
#         if y < 40:
#             c.showPage(); y = h - 40
#     c.save()

# def emit_doc(records):
#     # Plain-text .doc for Word compatibility
#     with open(os.path.join(OUT_DIR, "output.doc"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(r["text"] + "\n\n")

# def emit_docx(records):
#     doc = Document()
#     for r in records:
#         doc.add_paragraph(r["text"])
#     doc.save(os.path.join(OUT_DIR, "output.docx"))

# def emit_sql(records):
#     with open(os.path.join(OUT_DIR, "output.sql"), "w", encoding="utf-8") as f:
#         for r in records:
#             txt = r["text"].replace("'", "''")
#             var = r["variation"]
#             val = '1' if r.get("is_valid") else '0'
#             f.write(f"INSERT INTO sentences(text,variation,is_valid) VALUES('{txt}','{var}',{val});\n")

# def emit_diff(records):
#     with open(os.path.join(OUT_DIR, "output.diff"), "w", encoding="utf-8") as f:
#         f.write("--- original\n+++ mixed\n")
#         for r in records:
#             f.write(f"+ {r['text']}\n")
#         f.write("\n")

# # -----------------------------------------------------------------------------
# # 8. Paragraph emitter
# # -----------------------------------------------------------------------------
# def emit_paragraphs(records):
#     needed = PARA_COUNT * SENTS_PER_PARA
#     block  = records[:needed]
#     paras  = []
#     for i in range(PARA_COUNT):
#         chunk = block[i*SENTS_PER_PARA:(i+1)*SENTS_PER_PARA]
#         paras.append(" ".join(r["text"] for r in chunk))
#     with open(os.path.join(OUT_DIR, "output_paras.txt"), "w", encoding="utf-8") as f:
#         for p in paras:
#             f.write(p + "\n\n")

# # -----------------------------------------------------------------------------
# # 9. Main
# # -----------------------------------------------------------------------------
# if __name__ == "__main__":
#     count   = int(sys.argv[1]) if len(sys.argv)>1 else DEFAULT_COUNT
#     records = collect_records(count)
#     df      = pd.DataFrame(records)

#     emit_structured(df)
#     emit_pst(records)
#     emit_txt(records)
#     emit_log(records)
#     emit_html(records)
#     emit_pdf(records)
#     emit_doc(records)
#     emit_docx(records)
#     emit_sql(records)
#     emit_diff(records)
#     emit_paragraphs(records)

#     print(f"✅ Generated {count} records per module from {len(MODULES)} generators")
#     print(f"   → all outputs in folder: {OUT_DIR}")



# #!/usr/bin/env python3
# import os
# import sys
# import random
# import importlib
# import json
# import gzip
# import pandas as pd

# # PDF & DOCX
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# from docx import Document

# # -----------------------------------------------------------------------------
# # 1. Configuration
# # -----------------------------------------------------------------------------
# DEFAULT_COUNT   = 50    # records per module
# PARA_COUNT      = 5     # paragraphs
# SENTS_PER_PARA  = 5     # sentences each
# OUT_DIR         = "mixed_output"
# os.makedirs(OUT_DIR, exist_ok=True)

# # -----------------------------------------------------------------------------
# # 2. Discover & import generate_*.py modules
# # -----------------------------------------------------------------------------
# MODULES = []
# for fn in os.listdir("."):
#     if fn.startswith("generate_") and fn.endswith(".py"):
#         MODULES.append(importlib.import_module(fn[:-3]))

# # -----------------------------------------------------------------------------
# # 3. Locate each module’s *_variations() function
# # -----------------------------------------------------------------------------
# def get_generator(mod):
#     fns = [getattr(mod, n) for n in dir(mod)
#            if n.endswith("_variations") and callable(getattr(mod, n))]
#     if not fns:
#         raise AttributeError(f"No *_variations in {mod.__name__}")
#     if len(fns) == 1:
#         return fns[0]
#     base = mod.__name__.split("generate_")[-1]
#     for fn in fns:
#         if base in fn.__name__:
#             return fn
#     return fns[0]

# # -----------------------------------------------------------------------------
# # 4. Build & shuffle the pool
# # -----------------------------------------------------------------------------
# def collect_records(count_per_module):
#     pool = []
#     for mod in MODULES:
#         gen = get_generator(mod)
#         pool.extend(gen(count_per_module))
#     random.shuffle(pool)
#     return pool

# # -----------------------------------------------------------------------------
# # 5. Structured output
# # -----------------------------------------------------------------------------
# def emit_structured(df):
#     df.to_csv  (os.path.join(OUT_DIR, "output.csv"), index=False)
#     df.to_json (os.path.join(OUT_DIR, "output.json"), orient="records", indent=2)
#     df.to_excel(os.path.join(OUT_DIR, "output.xlsx"), index=False)
#     try:
#         df.to_excel(os.path.join(OUT_DIR, "output.xls"), index=False, engine="xlwt")
#     except Exception:
#         pass
#     try:
#         df.to_parquet(os.path.join(OUT_DIR, "output.parquet"), index=False)
#     except Exception:
#         pass

# # -----------------------------------------------------------------------------
# # 6. Compressed (.pst as GZIPped JSON)
# # -----------------------------------------------------------------------------
# def emit_compressed(records):
#     buf = json.dumps(records, ensure_ascii=False, indent=2).encode("utf-8")
#     with gzip.open(os.path.join(OUT_DIR, "output.pst"), "wb") as f:
#         f.write(buf)

# # -----------------------------------------------------------------------------
# # 7. Unstructured formats
# # -----------------------------------------------------------------------------
# def emit_txt(records):
#     with open(os.path.join(OUT_DIR, "output.txt"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(r["text"] + "\n\n")

# def emit_log(records):
#     with open(os.path.join(OUT_DIR, "output.log"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(f"{r['variation'].upper()}: {r['text']}\n")

# def emit_html(records):
#     with open(os.path.join(OUT_DIR, "output.html"), "w", encoding="utf-8") as f:
#         f.write("<!doctype html><html><body>\n")
#         for r in records:
#             f.write(f"  <p>{r['text']}</p>\n")
#         f.write("</body></html>")

# def emit_pdf(records):
#     c = canvas.Canvas(os.path.join(OUT_DIR, "output.pdf"), pagesize=letter)
#     w, h = letter; y = h - 40
#     for r in records:
#         c.drawString(40, y, r["text"])
#         y -= 14
#         if y < 40:
#             c.showPage(); y = h - 40
#     c.save()

# def emit_doc(records):
#     doc = Document()
#     for r in records:
#         doc.add_paragraph(r["text"])
#     doc.save(os.path.join(OUT_DIR, "output.doc"))

# def emit_docx(records):
#     doc = Document()
#     for r in records:
#         doc.add_paragraph(r["text"])
#     doc.save(os.path.join(OUT_DIR, "output.docx"))

# def emit_sql(records):
#     with open(os.path.join(OUT_DIR, "output.sql"), "w", encoding="utf-8") as f:
#         for r in records:
#             txt = r["text"].replace("'", "''")
#             var = r["variation"]
#             val = '1' if r.get("is_valid") else '0'
#             f.write(
#                 f"INSERT INTO sentences(text, variation, is_valid) "
#                 f"VALUES('{txt}', '{var}', {val});\n"
#             )

# def emit_diff(records):
#     with open(os.path.join(OUT_DIR, "output.diff"), "w", encoding="utf-8") as f:
#         f.write("--- original\n+++ mixed\n")
#         for r in records:
#             f.write(f"+ {r['text']}\n")
#         f.write("\n")

# # -----------------------------------------------------------------------------
# # 8. Paragraph emitter
# # -----------------------------------------------------------------------------
# def emit_paragraphs(records):
#     needed = PARA_COUNT * SENTS_PER_PARA
#     block  = records[:needed]
#     paras  = []
#     for i in range(PARA_COUNT):
#         chunk = block[i*SENTS_PER_PARA:(i+1)*SENTS_PER_PARA]
#         paras.append(" ".join(r["text"] for r in chunk))
#     with open(os.path.join(OUT_DIR, "output_paras.txt"), "w", encoding="utf-8") as f:
#         for p in paras:
#             f.write(p + "\n\n")

# # -----------------------------------------------------------------------------
# # 9. Main
# # -----------------------------------------------------------------------------
# if __name__ == "__main__":
#     count   = int(sys.argv[1]) if len(sys.argv)>1 else DEFAULT_COUNT
#     records = collect_records(count)
#     df      = pd.DataFrame(records)

#     emit_structured(df)
#     emit_compressed(records)
#     emit_txt(records)
#     emit_log(records)
#     emit_html(records)
#     emit_pdf(records)
#     emit_doc(records)
#     emit_docx(records)
#     emit_sql(records)
#     emit_diff(records)
#     emit_paragraphs(records)

#     print(f"✅ Generated {count} records per module from {len(MODULES)} generators")
#     print(f"   → all outputs in folder: {OUT_DIR}")



# #!/usr/bin/env python3
# import os
# import sys
# import random
# import importlib
# import json
# import pandas as pd

# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# from docx import Document

# # -----------------------------------------------------------------------------
# # 1. Configuration
# # -----------------------------------------------------------------------------
# DEFAULT_COUNT   = 50    # records per module
# PARA_COUNT      = 5     # paragraphs
# SENTS_PER_PARA  = 5     # sentences each
# OUT_DIR         = "mixed_output"
# os.makedirs(OUT_DIR, exist_ok=True)

# # -----------------------------------------------------------------------------
# # 2. Discover & import generate_*.py modules
# # -----------------------------------------------------------------------------
# MODULES = []
# for fn in os.listdir("."):
#     if fn.startswith("generate_") and fn.endswith(".py"):
#         MODULES.append(importlib.import_module(fn[:-3]))

# # -----------------------------------------------------------------------------
# # 3. Locate each module’s *_variations() function
# # -----------------------------------------------------------------------------
# def get_generator(mod):
#     funcs = [getattr(mod, name) for name in dir(mod)
#              if name.endswith("_variations") and callable(getattr(mod, name))]
#     if not funcs:
#         raise AttributeError(f"No `*_variations` in {mod.__name__}")
#     if len(funcs) == 1:
#         return funcs[0]
#     # tie‑break by module base name
#     base = mod.__name__.split("generate_")[-1]
#     for fn in funcs:
#         if base in fn.__name__:
#             return fn
#     return funcs[0]

# # -----------------------------------------------------------------------------
# # 4. Build & shuffle the pool
# # -----------------------------------------------------------------------------
# def collect_records(count_per_module):
#     pool = []
#     for mod in MODULES:
#         gen = get_generator(mod)
#         pool.extend(gen(count_per_module))
#     random.shuffle(pool)
#     return pool

# # -----------------------------------------------------------------------------
# # 5. Structured output
# # -----------------------------------------------------------------------------
# def emit_structured(df):
#     # CSV
#     df.to_csv(os.path.join(OUT_DIR, "output.csv"), index=False)
#     # JSON without nulls
#     records = []
#     for rec in df.to_dict(orient="records"):
#         clean = {k: v for k, v in rec.items() if v is not None and v != ""}
#         records.append(clean)
#     with open(os.path.join(OUT_DIR, "output.json"), "w", encoding="utf-8") as f:
#         json.dump(records, f, ensure_ascii=False, indent=2)
#     # XLSX
#     df.to_excel(os.path.join(OUT_DIR, "output.xlsx"), index=False)
#     # XLS (optional)
#     try:
#         import xlwt
#         df.to_excel(os.path.join(OUT_DIR, "output.xls"), index=False, engine="xlwt")
#     except ImportError:
#         pass
#     # Parquet
#     try:
#         df.to_parquet(os.path.join(OUT_DIR, "output.parquet"), index=False)
#     except Exception:
#         pass
#     # PST = JSON dump (no nulls)
#     with open(os.path.join(OUT_DIR, "output.pst"), "w", encoding="utf-8") as f:
#         json.dump(records, f, ensure_ascii=False, indent=2)

# # -----------------------------------------------------------------------------
# # 6. Unstructured
# # -----------------------------------------------------------------------------
# def emit_txt(records):
#     with open(os.path.join(OUT_DIR, "output.txt"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(r["text"] + "\n\n")

# def emit_log(records):
#     with open(os.path.join(OUT_DIR, "output.log"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(f"{r['variation'].upper()}: {r['text']}\n")

# def emit_html(records):
#     with open(os.path.join(OUT_DIR, "output.html"), "w", encoding="utf-8") as f:
#         f.write("<!doctype html><html><body>\n")
#         for r in records:
#             f.write(f"  <p>{r['text']}</p>\n")
#         f.write("</body></html>")

# def emit_pdf(records):
#     c = canvas.Canvas(os.path.join(OUT_DIR, "output.pdf"), pagesize=letter)
#     w, h = letter
#     y = h - 40
#     for r in records:
#         c.drawString(40, y, r["text"])
#         y -= 14
#         if y < 40:
#             c.showPage()
#             y = h - 40
#     c.save()

# def emit_doc(records):
#     # Rich‐text .doc via python‑docx
#     doc = Document()
#     for r in records:
#         doc.add_paragraph(r["text"])
#     doc.save(os.path.join(OUT_DIR, "output.doc"))

# def emit_docx(records):
#     emit_doc(records)  # same as .doc for simplicity
#     os.replace(os.path.join(OUT_DIR, "output.doc"),
#                os.path.join(OUT_DIR, "output.docx"))

# def emit_sql(records):
#     with open(os.path.join(OUT_DIR, "output.sql"), "w", encoding="utf-8") as f:
#         for r in records:
#             txt = r["text"].replace("'", "''")
#             var = r["variation"]
#             val = '1' if r.get("is_valid") else '0'
#             f.write(f"INSERT INTO sentences(text,variation,is_valid) "
#                     f"VALUES('{txt}','{var}',{val});\n")

# def emit_diff(records):
#     with open(os.path.join(OUT_DIR, "output.diff"), "w", encoding="utf-8") as f:
#         f.write("--- original\n+++ mixed\n")
#         for r in records:
#             f.write(f"+ {r['text']}\n")
#         f.write("\n")

# def emit_paras(records):
#     needed = PARA_COUNT * SENTS_PER_PARA
#     block = records[:needed]
#     paras = []
#     for i in range(PARA_COUNT):
#         chunk = block[i*SENTS_PER_PARA:(i+1)*SENTS_PER_PARA]
#         paras.append(" ".join(r["text"] for r in chunk))
#     with open(os.path.join(OUT_DIR, "output_paras.txt"), "w", encoding="utf-8") as f:
#         for p in paras:
#             f.write(p + "\n\n")

# # -----------------------------------------------------------------------------
# # 7. Main
# # -----------------------------------------------------------------------------
# if __name__ == "__main__":
#     count = int(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_COUNT
#     records = collect_records(count)
#     df = pd.DataFrame(records)

#     emit_structured(df)
#     emit_txt(records)
#     emit_log(records)
#     emit_html(records)
#     emit_pdf(records)
#     emit_doc(records)
#     emit_docx(records)
#     emit_sql(records)
#     emit_diff(records)
#     emit_paras(records)

#     print(f"✅ Generated {count} records per module from {len(MODULES)} generators")
#     print(f"   → all outputs in folder: {OUT_DIR}")



#!/usr/bin/env python3
# import os
# import sys
# import random
# import importlib
# import json
# import pandas as pd

# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# from docx import Document

# # -----------------------------------------------------------------------------
# # 1. Configuration
# # -----------------------------------------------------------------------------
# DEFAULT_COUNT   = 50    # records per module
# PARA_COUNT      = 5     # paragraphs
# SENTS_PER_PARA  = 5     # sentences each
# OUT_DIR         = "mixed_output"
# os.makedirs(OUT_DIR, exist_ok=True)

# # -----------------------------------------------------------------------------
# # 2. Discover & import generate_*.py modules
# # -----------------------------------------------------------------------------
# MODULES = []
# for fn in os.listdir("."):
#     if fn.startswith("generate_") and fn.endswith(".py"):
#         MODULES.append(importlib.import_module(fn[:-3]))

# # -----------------------------------------------------------------------------
# # 3. Locate each module’s *_variations() function
# # -----------------------------------------------------------------------------
# def get_generator(mod):
#     fns = [
#         getattr(mod, n) for n in dir(mod)
#         if n.endswith("_variations") and callable(getattr(mod, n))
#     ]
#     if not fns:
#         raise AttributeError(f"No `*_variations` in {mod.__name__}")
#     if len(fns) == 1:
#         return fns[0]
#     base = mod.__name__.split("generate_")[-1]
#     for fn in fns:
#         if base in fn.__name__:
#             return fn
#     return fns[0]

# # -----------------------------------------------------------------------------
# # 4. Build & shuffle the record pool
# # -----------------------------------------------------------------------------
# def collect_records(count_per_module):
#     pool = []
#     for mod in MODULES:
#         gen = get_generator(mod)
#         pool.extend(gen(count_per_module))
#     random.shuffle(pool)
#     return pool

# # -----------------------------------------------------------------------------
# # 5. Structured output
# # -----------------------------------------------------------------------------
# def emit_structured(df):
#     # 5.1 CSV
#     df.to_csv(os.path.join(OUT_DIR, "output.csv"), index=False)

#     # 5.2 JSON (no nulls): only include keys with non‑empty values
#     records = []
#     for rec in df.to_dict(orient="records"):
#         clean = {k: v for k, v in rec.items() if v not in (None, "", [], {})}
#         records.append(clean)
#     with open(os.path.join(OUT_DIR, "output.json"), "w", encoding="utf-8") as f:
#         json.dump(records, f, ensure_ascii=False, indent=2)

#     # 5.3 XLSX
#     df.to_excel(os.path.join(OUT_DIR, "output.xlsx"), index=False)

#     # 5.4 XLS (optional) – catch any errors if xlwt is missing
#     try:
#         df.to_excel(os.path.join(OUT_DIR, "output.xls"), index=False, engine="xlwt")
#     except Exception:
#         pass

#     # 5.5 Parquet
#     try:
#         df.to_parquet(os.path.join(OUT_DIR, "output.parquet"), index=False)
#     except Exception:
#         pass

#     # 5.6 PST (.pst = JSON dump with same non‑null records)
#     with open(os.path.join(OUT_DIR, "output.pst"), "w", encoding="utf-8") as f:
#         json.dump(records, f, ensure_ascii=False, indent=2)

# # -----------------------------------------------------------------------------
# # 6. Unstructured formats
# # -----------------------------------------------------------------------------
# def emit_txt(records):
#     with open(os.path.join(OUT_DIR, "output.txt"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(r["text"] + "\n\n")

# def emit_log(records):
#     with open(os.path.join(OUT_DIR, "output.log"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(f"{r['variation'].upper()}: {r['text']}\n")

# def emit_html(records):
#     with open(os.path.join(OUT_DIR, "output.html"), "w", encoding="utf-8") as f:
#         f.write("<!doctype html>\n<html lang=\"en\"><body>\n")
#         for r in records:
#             f.write(f"  <p>{r['text']}</p>\n")
#         f.write("</body></html>")

# def emit_pdf(records):
#     c = canvas.Canvas(os.path.join(OUT_DIR, "output.pdf"), pagesize=letter)
#     w, h = letter
#     y = h - 40
#     for r in records:
#         c.drawString(40, y, r["text"])
#         y -= 14
#         if y < 40:
#             c.showPage()
#             y = h - 40
#     c.save()

# def emit_docx(records):
#     doc = Document()
#     for r in records:
#         doc.add_paragraph(r["text"])
#     doc.save(os.path.join(OUT_DIR, "output.docx"))

# def emit_sql(records):
#     with open(os.path.join(OUT_DIR, "output.sql"), "w", encoding="utf-8") as f:
#         for r in records:
#             txt = r["text"].replace("'", "''")
#             var = r["variation"]
#             val = '1' if r.get("is_valid") else '0'
#             f.write(f"INSERT INTO sentences(text,variation,is_valid) VALUES('{txt}','{var}',{val});\n")

# def emit_diff(records):
#     with open(os.path.join(OUT_DIR, "output.diff"), "w", encoding="utf-8") as f:
#         f.write("--- original\n+++ mixed\n")
#         for r in records:
#             f.write(f"+ {r['text']}\n")
#         f.write("\n")

# # -----------------------------------------------------------------------------
# # 7. Paragraph emitter
# # -----------------------------------------------------------------------------
# def emit_paragraphs(records):
#     needed = PARA_COUNT * SENTS_PER_PARA
#     block  = records[:needed]
#     paras  = [
#         " ".join(r["text"] for r in block[i*SENTS_PER_PARA:(i+1)*SENTS_PER_PARA])
#         for i in range(PARA_COUNT)
#     ]
#     with open(os.path.join(OUT_DIR, "output_paras.txt"), "w", encoding="utf-8") as f:
#         for p in paras:
#             f.write(p + "\n\n")

# # -----------------------------------------------------------------------------
# # 8. Main
# # -----------------------------------------------------------------------------
# if __name__ == "__main__":
#     count   = int(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_COUNT
#     records = collect_records(count)
#     df      = pd.DataFrame(records)

#     emit_structured(df)
#     emit_txt(records)
#     emit_log(records)
#     emit_html(records)
#     emit_pdf(records)
#     emit_docx(records)
#     emit_sql(records)
#     emit_diff(records)
#     emit_paragraphs(records)

#     print(f"✅ Generated {count} records per module from {len(MODULES)} generators")
#     print(f"   → all outputs in folder: {OUT_DIR}")



#!/usr/bin/env python3
# import os
# import sys
# import random
# import importlib
# import json
# import pandas as pd

# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# from docx import Document

# # -----------------------------------------------------------------------------
# # 1. Configuration
# # -----------------------------------------------------------------------------
# DEFAULT_COUNT   = 50    # records per module
# PARA_COUNT      = 5     # paragraphs
# SENTS_PER_PARA  = 5     # sentences each
# OUT_DIR         = "mixed_output"
# os.makedirs(OUT_DIR, exist_ok=True)

# # -----------------------------------------------------------------------------
# # 2. Discover & import generate_*.py modules
# # -----------------------------------------------------------------------------
# MODULES = []
# for fn in os.listdir("."):
#     if fn.startswith("generate_") and fn.endswith(".py"):
#         MODULES.append(importlib.import_module(fn[:-3]))

# # -----------------------------------------------------------------------------
# # 3. Locate each module’s *_variations() function
# # -----------------------------------------------------------------------------
# def get_generator(mod):
#     fns = [getattr(mod, n) for n in dir(mod)
#            if n.endswith("_variations") and callable(getattr(mod, n))]
#     if not fns:
#         raise AttributeError(f"No *_variations in {mod.__name__}")
#     if len(fns)==1:
#         return fns[0]
#     # disambiguate by module name
#     base = mod.__name__.split("generate_")[-1]
#     for fn in fns:
#         if base in fn.__name__:
#             return fn
#     return fns[0]

# # -----------------------------------------------------------------------------
# # 4. Build & shuffle the pool
# # -----------------------------------------------------------------------------
# def collect_records(count_per_module):
#     pool = []
#     for mod in MODULES:
#         gen = get_generator(mod)
#         pool.extend(gen(count_per_module))
#     random.shuffle(pool)
#     return pool

# # -----------------------------------------------------------------------------
# # 5. Structured output
# # -----------------------------------------------------------------------------
# def emit_structured(df):
#     #  CSV & XLSX & Parquet
#     df.to_csv  ( os.path.join(OUT_DIR, "output.csv"), index=False )
#     df.to_excel( os.path.join(OUT_DIR, "output.xlsx"), index=False )
#     try:
#         df.to_parquet(os.path.join(OUT_DIR, "output.parquet"), index=False)
#     except Exception:
#         pass

#     # JSON – skip nulls per record
#     out_json = []
#     for rec in df.to_dict(orient="records"):
#         clean = {k: v for k, v in rec.items() if v is not None}
#         out_json.append(clean)
#     with open(os.path.join(OUT_DIR, "output.json"), "w", encoding="utf-8") as f:
#         json.dump(out_json, f, ensure_ascii=False, indent=2)

#     # PST – same as JSON
#     with open(os.path.join(OUT_DIR, "output.pst"), "w", encoding="utf-8") as f:
#         json.dump(out_json, f, ensure_ascii=False, indent=2)

# # -----------------------------------------------------------------------------
# # 6. Unstructured
# # -----------------------------------------------------------------------------
# def emit_txt(records):
#     with open(os.path.join(OUT_DIR, "output.txt"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(r["text"] + "\n\n")

# def emit_log(records):
#     with open(os.path.join(OUT_DIR, "output.log"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(f"{r['variation'].upper()}: {r['text']}\n")

# def emit_sql(records):
#     with open(os.path.join(OUT_DIR, "output.sql"), "w", encoding="utf-8") as f:
#         for r in records:
#             txt = r["text"].replace("'", "''")
#             var = r["variation"]
#             val = '1' if r.get("is_valid") else '0'
#             f.write(
#                 f"INSERT INTO sentences(text,variation,is_valid) "
#                 f"VALUES('{txt}','{var}',{val});\n"
#             )

# def emit_diff(records):
#     with open(os.path.join(OUT_DIR, "output.diff"), "w", encoding="utf-8") as f:
#         f.write("--- original\n+++ mixed\n")
#         for r in records:
#             f.write(f"+ {r['text']}\n")
#         f.write("\n")

# def emit_html(records):
#     with open(os.path.join(OUT_DIR, "output.html"), "w", encoding="utf-8") as f:
#         f.write("<!doctype html>\n<html><body>\n")
#         for r in records:
#             f.write(f"  <p>{r['text']}</p>\n")
#         f.write("</body></html>")

# def emit_pdf(records):
#     c = canvas.Canvas(os.path.join(OUT_DIR, "output.pdf"), pagesize=letter)
#     w,h = letter
#     y = h-40
#     for r in records:
#         c.drawString(40, y, r["text"])
#         y -= 14
#         if y < 40:
#             c.showPage()
#             y = h-40
#     c.save()

# def emit_doc(records):
#     doc = Document()
#     for r in records:
#         doc.add_paragraph(r["text"])
#     doc.save(os.path.join(OUT_DIR, "output.doc"))

# def emit_docx(records):
#     doc = Document()
#     for r in records:
#         doc.add_paragraph(r["text"])
#     doc.save(os.path.join(OUT_DIR, "output.docx"))

# def emit_paragraphs(records):
#     needed = PARA_COUNT * SENTS_PER_PARA
#     block  = records[:needed]
#     paras  = []
#     for i in range(PARA_COUNT):
#         chunk = block[i*SENTS_PER_PARA:(i+1)*SENTS_PER_PARA]
#         paras.append(" ".join(r["text"] for r in chunk))
#     with open(os.path.join(OUT_DIR, "output_paras.txt"), "w", encoding="utf-8") as f:
#         for p in paras:
#             f.write(p + "\n\n")

# # -----------------------------------------------------------------------------
# # 7. Main
# # -----------------------------------------------------------------------------
# if __name__ == "__main__":
#     count   = int(sys.argv[1]) if len(sys.argv)>1 else DEFAULT_COUNT
#     records = collect_records(count)
#     df      = pd.DataFrame(records)

#     emit_structured(df)
#     emit_txt(records)
#     emit_log(records)
#     emit_sql(records)
#     emit_diff(records)
#     emit_html(records)
#     emit_pdf(records)
#     emit_doc(records)
#     emit_docx(records)
#     emit_paragraphs(records)

#     print(f"✅ Generated {count} records per module from {len(MODULES)} generators")
#     print(f"   → all outputs in folder: {OUT_DIR}")




#!/usr/bin/env python3
# import os
# import sys
# import random
# import importlib
# import json
# import math

# import pandas as pd
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# from docx import Document

# # -----------------------------------------------------------------------------
# # 1. Configuration
# # -----------------------------------------------------------------------------
# DEFAULT_COUNT   = 50    # records per module
# PARA_COUNT      = 5     # paragraphs
# SENTS_PER_PARA  = 5     # sentences each
# OUT_DIR         = "mixed_output"
# os.makedirs(OUT_DIR, exist_ok=True)

# # -----------------------------------------------------------------------------
# # 2. Discover & import generate_*.py modules
# # -----------------------------------------------------------------------------
# MODULES = []
# for fn in os.listdir("."):
#     if fn.startswith("generate_") and fn.endswith(".py"):
#         MODULES.append(importlib.import_module(fn[:-3]))

# # -----------------------------------------------------------------------------
# # 3. Locate each module’s *_variations() function
# # -----------------------------------------------------------------------------
# def get_generator(mod):
#     funcs = [
#         getattr(mod, name)
#         for name in dir(mod)
#         if name.endswith("_variations") and callable(getattr(mod, name))
#     ]
#     if not funcs:
#         raise AttributeError(f"No `*_variations` in {mod.__name__}")
#     if len(funcs) == 1:
#         return funcs[0]
#     base = mod.__name__.split("generate_")[-1]
#     for fn in funcs:
#         if base in fn.__name__:
#             return fn
#     return funcs[0]

# # -----------------------------------------------------------------------------
# # 4. Build & shuffle the pool
# # -----------------------------------------------------------------------------
# def collect_records(count_per_module):
#     pool = []
#     for mod in MODULES:
#         gen = get_generator(mod)
#         pool.extend(gen(count_per_module))
#     random.shuffle(pool)
#     return pool

# # -----------------------------------------------------------------------------
# # 5. Structured output
# # -----------------------------------------------------------------------------
# def emit_structured(df):
#     # (a) CSV & XLSX
#     df.to_csv  ( os.path.join(OUT_DIR, "output.csv"), index=False )
#     df.to_excel( os.path.join(OUT_DIR, "output.xlsx"), index=False )

#     # (b) try Parquet
#     try:
#         df.to_parquet(os.path.join(OUT_DIR, "output.parquet"), index=False)
#     except Exception:
#         pass

#     # (c) Build a “clean” list of dicts without None/NaN values
#     clean_recs = []
#     for rec in df.to_dict(orient="records"):
#         clean = {}
#         for k,v in rec.items():
#             if v is None:
#                 continue
#             if isinstance(v, float) and math.isnan(v):
#                 continue
#             clean[k] = v
#         clean_recs.append(clean)

#     # (d) JSON
#     with open(os.path.join(OUT_DIR, "output.json"), "w", encoding="utf-8") as f:
#         json.dump(clean_recs, f, ensure_ascii=False, indent=2)

#     # (e) PST (same JSON dump)
#     with open(os.path.join(OUT_DIR, "output.pst"), "w", encoding="utf-8") as f:
#         json.dump(clean_recs, f, ensure_ascii=False, indent=2)

#     # (f) optional old-school XLS (skip if no xlwt)
#     try:
#         df.to_excel(os.path.join(OUT_DIR, "output.xls"), index=False, engine="xlwt")
#     except Exception:
#         pass

# # -----------------------------------------------------------------------------
# # 6. Unstructured + other formats
# # -----------------------------------------------------------------------------
# def emit_txt(records):
#     with open(os.path.join(OUT_DIR, "output.txt"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(r["text"] + "\n\n")

# def emit_log(records):
#     with open(os.path.join(OUT_DIR, "output.log"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(f"{r['variation'].upper()}: {r['text']}\n")

# def emit_html(records):
#     # now bona fide HTML
#     with open(os.path.join(OUT_DIR, "output.html"), "w", encoding="utf-8") as f:
#         f.write("<!doctype html>\n<html><head><meta charset='utf-8'><title>Mixed Output</title></head><body>\n")
#         for r in records:
#             f.write(f"  <p>{r['text']}</p>\n")
#         f.write("</body></html>")

# def emit_pdf(records):
#     c = canvas.Canvas(os.path.join(OUT_DIR, "output.pdf"), pagesize=letter)
#     w, h = letter
#     y = h - 40
#     for r in records:
#         c.drawString(40, y, r["text"])
#         y -= 14
#         if y < 40:
#             c.showPage()
#             y = h - 40
#     c.save()

# def emit_docx(records):
#     doc = Document()
#     for r in records:
#         doc.add_paragraph(r["text"])
#     doc.save(os.path.join(OUT_DIR, "output.docx"))

# def emit_sql(records):
#     with open(os.path.join(OUT_DIR, "output.sql"), "w", encoding="utf-8") as f:
#         for r in records:
#             txt = r["text"].replace("'", "''")
#             var = r["variation"]
#             val = '1' if r.get("is_valid") else '0'
#             f.write(f"INSERT INTO sentences(text,variation,is_valid) VALUES('{txt}','{var}',{val});\n")

# def emit_diff(records):
#     with open(os.path.join(OUT_DIR, "output.diff"), "w", encoding="utf-8") as f:
#         f.write("--- original\n+++ mixed\n")
#         for r in records:
#             f.write(f"+ {r['text']}\n")
#         f.write("\n")

# # -----------------------------------------------------------------------------
# # 7. Paragraph emitter
# # -----------------------------------------------------------------------------
# def emit_paragraphs(records):
#     needed = PARA_COUNT * SENTS_PER_PARA
#     block  = records[:needed]
#     paras  = []
#     for i in range(PARA_COUNT):
#         chunk = block[i*SENTS_PER_PARA:(i+1)*SENTS_PER_PARA]
#         paras.append(" ".join(r["text"] for r in chunk))
#     with open(os.path.join(OUT_DIR, "output_paras.txt"), "w", encoding="utf-8") as f:
#         for p in paras:
#             f.write(p + "\n\n")

# # -----------------------------------------------------------------------------
# # 8. Main
# # -----------------------------------------------------------------------------
# if __name__ == "__main__":
#     count   = int(sys.argv[1]) if len(sys.argv)>1 else DEFAULT_COUNT
#     records = collect_records(count)
#     df      = pd.DataFrame(records)

#     emit_structured(df)
#     emit_txt(records)
#     emit_log(records)
#     emit_html(records)
#     emit_pdf(records)
#     emit_docx(records)
#     emit_sql(records)
#     emit_diff(records)
#     emit_paragraphs(records)

#     print(f"✅ Generated {count} records per module from {len(MODULES)} generators")
#     print(f"   → all outputs in folder: {OUT_DIR}")



# #!/usr/bin/env python3
# import os, sys, random, importlib, json
# import pandas as pd
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# from docx import Document

# # -----------------------------------------------------------------------------
# # 1. Configuration
# # -----------------------------------------------------------------------------
# DEFAULT_COUNT   = 50    # records per module
# PARA_COUNT      = 5     # paragraphs
# SENTS_PER_PARA  = 5     # sentences each
# OUT_DIR         = "mixed_output"
# os.makedirs(OUT_DIR, exist_ok=True)

# # -----------------------------------------------------------------------------
# # 2. Discover & import generate_*.py modules
# # -----------------------------------------------------------------------------
# MODULES = []
# for fn in os.listdir("."):
#     if fn.startswith("generate_") and fn.endswith(".py"):
#         MODULES.append(importlib.import_module(fn[:-3]))

# # -----------------------------------------------------------------------------
# # 3. Locate each module’s *_variations() function
# # -----------------------------------------------------------------------------
# def get_generator(mod):
#     fns = [getattr(mod, n) for n in dir(mod)
#            if n.endswith("_variations") and callable(getattr(mod, n))]
#     if not fns:
#         raise AttributeError(f"No `*_variations` in {mod.__name__}")
#     if len(fns) == 1:
#         return fns[0]
#     base = mod.__name__.split("generate_")[-1]
#     for fn in fns:
#         if base in fn.__name__:
#             return fn
#     return fns[0]

# # -----------------------------------------------------------------------------
# # 4. Build & shuffle the pool
# # -----------------------------------------------------------------------------
# def collect_records(count_per_module):
#     pool = []
#     for mod in MODULES:
#         gen = get_generator(mod)
#         pool.extend(gen(count_per_module))
#     random.shuffle(pool)
#     return pool

# # -----------------------------------------------------------------------------
# # 5. Structured outputs
# # -----------------------------------------------------------------------------
# def emit_structured(df):
#     # CSV
#     df.to_csv(os.path.join(OUT_DIR, "output.csv"), index=False)

#     # JSON (replace NaN with None → null)
#     clean = df.where(pd.notnull(df), None)
#     with open(os.path.join(OUT_DIR, "output.json"), "w", encoding="utf-8") as f:
#         json.dump(clean.to_dict(orient="records"), f, ensure_ascii=False, indent=2)

#     # XLSX
#     df.to_excel(os.path.join(OUT_DIR, "output.xlsx"), index=False)

#     # Parquet
#     try:
#         df.to_parquet(os.path.join(OUT_DIR, "output.parquet"), index=False)
#     except Exception:
#         pass

# # -----------------------------------------------------------------------------
# # 6. PST as SQL dump
# # -----------------------------------------------------------------------------
# def emit_pst_sql(records):
#     with open(os.path.join(OUT_DIR, "output.pst"), "w", encoding="utf-8") as f:
#         for r in records:
#             txt = r["text"].replace("'", "''")
#             var = r["variation"]
#             val = '1' if r.get("is_valid") else '0'
#             f.write(f"INSERT INTO sentences(text,variation,is_valid) VALUES('{txt}','{var}',{val});\n")

# # -----------------------------------------------------------------------------
# # 7. Unstructured formats
# # -----------------------------------------------------------------------------
# def emit_txt(records):
#     with open(os.path.join(OUT_DIR, "output.txt"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(r["text"] + "\n\n")

# def emit_log(records):
#     with open(os.path.join(OUT_DIR, "output.log"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(r["text"] + "\n")

# def emit_diff(records):
#     with open(os.path.join(OUT_DIR, "output.diff"), "w", encoding="utf-8") as f:
#         f.write("--- original\n+++ mixed\n")
#         for r in records:
#             f.write(f"+ {r['text']}\n")
#         f.write("\n")

# def emit_html(records):
#     # Real HTML with <p> wrappers
#     with open(os.path.join(OUT_DIR, "output.html"), "w", encoding="utf-8") as f:
#         f.write("<!doctype html>\n<html><head><meta charset='utf-8'><title>Mixed Output</title></head><body>\n")
#         for r in records:
#             f.write(f"  <p>{r['text']}</p>\n")
#         f.write("</body></html>")

# # -----------------------------------------------------------------------------
# # 8. Paragraph‑text
# # -----------------------------------------------------------------------------
# def emit_paragraphs(records):
#     needed = PARA_COUNT * SENTS_PER_PARA
#     block  = records[:needed]
#     paras  = []
#     for i in range(PARA_COUNT):
#         chunk = block[i*SENTS_PER_PARA:(i+1)*SENTS_PER_PARA]
#         paras.append(" ".join(r["text"] for r in chunk))
#     with open(os.path.join(OUT_DIR, "output_paras.txt"), "w", encoding="utf-8") as f:
#         for p in paras:
#             f.write(p + "\n\n")

# # -----------------------------------------------------------------------------
# # 9. PDF, DOCX, SQL
# # -----------------------------------------------------------------------------
# def emit_pdf(records):
#     c = canvas.Canvas(os.path.join(OUT_DIR, "output.pdf"), pagesize=letter)
#     w,h = letter; y = h-40
#     for r in records:
#         c.drawString(40, y, r["text"])
#         y -= 14
#         if y < 40:
#             c.showPage(); y = h-40
#     c.save()

# def emit_docx(records):
#     doc = Document()
#     for r in records:
#         doc.add_paragraph(r["text"])
#     doc.save(os.path.join(OUT_DIR, "output.docx"))

# def emit_sql(records):
#     with open(os.path.join(OUT_DIR, "output.sql"), "w", encoding="utf-8") as f:
#         for r in records:
#             txt = r["text"].replace("'", "''")
#             var = r["variation"]
#             val = '1' if r.get("is_valid") else '0'
#             f.write(f"INSERT INTO sentences(text,variation,is_valid) VALUES('{txt}','{var}',{val});\n")

# # -----------------------------------------------------------------------------
# # 10. Main
# # -----------------------------------------------------------------------------
# if __name__ == "__main__":
#     count   = int(sys.argv[1]) if len(sys.argv)>1 else DEFAULT_COUNT
#     records = collect_records(count)
#     df      = pd.DataFrame(records)

#     emit_structured(df)
#     emit_pst_sql(records)
#     emit_txt(records)
#     emit_log(records)
#     emit_diff(records)
#     emit_html(records)
#     emit_paragraphs(records)
#     emit_pdf(records)
#     emit_docx(records)
#     emit_sql(records)

#     print(f"✅ Generated {count} records per module from {len(MODULES)} generators")
#     print(f"   → all outputs in folder: {OUT_DIR}")



# #!/usr/bin/env python3
# import os, sys, random, importlib, json
# import pandas as pd

# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# from docx import Document

# # -----------------------------------------------------------------------------
# # 1. Configuration
# # -----------------------------------------------------------------------------
# DEFAULT_COUNT   = 50
# PARA_COUNT      = 5
# SENTS_PER_PARA  = 5
# OUT_DIR         = "mixed_output"
# os.makedirs(OUT_DIR, exist_ok=True)

# # -----------------------------------------------------------------------------
# # 2. Discover & import your generate_*.py modules
# # -----------------------------------------------------------------------------
# MODULES = []
# for fn in os.listdir("."):
#     if fn.startswith("generate_") and fn.endswith(".py"):
#         MODULES.append(importlib.import_module(fn[:-3]))

# # -----------------------------------------------------------------------------
# # 3. Locate each module’s *_variations() function
# # -----------------------------------------------------------------------------
# def get_generator(mod):
#     funcs = [getattr(mod, n) for n in dir(mod)
#              if n.endswith("_variations") and callable(getattr(mod, n))]
#     if not funcs:
#         raise AttributeError(f"No `*_variations` in {mod.__name__}")
#     if len(funcs) == 1:
#         return funcs[0]
#     base = mod.__name__.split("generate_")[-1]
#     for fn in funcs:
#         if base in fn.__name__:
#             return fn
#     return funcs[0]

# # -----------------------------------------------------------------------------
# # 4. Collect & shuffle
# # -----------------------------------------------------------------------------
# def collect_records(count):
#     pool = []
#     for m in MODULES:
#         gen = get_generator(m)
#         pool.extend(gen(count))
#     random.shuffle(pool)
#     return pool

# # -----------------------------------------------------------------------------
# # 5. Structured outputs
# # -----------------------------------------------------------------------------
# def emit_csv(df):
#     df.to_csv(os.path.join(OUT_DIR, "output.csv"), index=False)

# def emit_xlsx(df):
#     df.to_excel(os.path.join(OUT_DIR, "output.xlsx"), index=False)

# def emit_parquet(df):
#     try:
#         df.to_parquet(os.path.join(OUT_DIR, "output.parquet"), index=False)
#     except Exception:
#         pass

# def emit_json(records):
#     # filter out missing fields so no nulls appear
#     out = []
#     for r in records:
#         obj = {k: v for k, v in r.items() if v is not None}
#         out.append(obj)
#     with open(os.path.join(OUT_DIR, "output.json"), "w", encoding="utf-8") as f:
#         json.dump(out, f, ensure_ascii=False, indent=2)

# # -----------------------------------------------------------------------------
# # 6. PST as SQL
# # -----------------------------------------------------------------------------
# def emit_pst(records):
#     with open(os.path.join(OUT_DIR, "output.pst"), "w", encoding="utf-8") as f:
#         for r in records:
#             txt = r["text"].replace("'", "''")
#             var = r["variation"]
#             valid = '1' if r.get("is_valid") else '0'
#             f.write(f"INSERT INTO sentences(text,variation,is_valid) VALUES('{txt}','{var}',{valid});\n")

# # -----------------------------------------------------------------------------
# # 7. Unstructured
# # -----------------------------------------------------------------------------
# def emit_txt(records):
#     with open(os.path.join(OUT_DIR, "output.txt"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(r["text"] + "\n\n")

# def emit_log(records):
#     with open(os.path.join(OUT_DIR, "output.log"), "w", encoding="utf-8") as f:
#         for r in records:
#             f.write(r["text"] + "\n")

# def emit_diff(records):
#     with open(os.path.join(OUT_DIR, "output.diff"), "w", encoding="utf-8") as f:
#         f.write("--- original\n+++ mixed\n")
#         for r in records:
#             f.write(f"+ {r['text']}\n")
#         f.write("\n")

# def emit_html(records):
#     with open(os.path.join(OUT_DIR, "output.html"), "w", encoding="utf-8") as f:
#         f.write("<!doctype html>\n<html><head><meta charset='utf-8'><title>Mixed Output</title></head><body>\n")
#         for r in records:
#             f.write(f"  <p>{r['text']}</p>\n")
#         f.write("</body></html>")

# # -----------------------------------------------------------------------------
# # 8. Paragraphs
# # -----------------------------------------------------------------------------
# def emit_paragraphs(records):
#     needed = PARA_COUNT * SENTS_PER_PARA
#     chunk  = records[:needed]
#     with open(os.path.join(OUT_DIR, "output_paras.txt"), "w", encoding="utf-8") as f:
#         for i in range(PARA_COUNT):
#             para = " ".join(r["text"] for r in chunk[i*SENTS_PER_PARA:(i+1)*SENTS_PER_PARA])
#             f.write(para + "\n\n")

# # -----------------------------------------------------------------------------
# # 9. PDF, DOCX, SQL
# # -----------------------------------------------------------------------------
# def emit_pdf(records):
#     c = canvas.Canvas(os.path.join(OUT_DIR, "output.pdf"), pagesize=letter)
#     w, h = letter; y = h-40
#     for r in records:
#         c.drawString(40, y, r["text"])
#         y -= 14
#         if y < 40:
#             c.showPage(); y = h-40
#     c.save()

# def emit_docx(records):
#     doc = Document()
#     for r in records:
#         doc.add_paragraph(r["text"])
#     doc.save(os.path.join(OUT_DIR, "output.docx"))

# def emit_sql(records):
#     with open(os.path.join(OUT_DIR, "output.sql"), "w", encoding="utf-8") as f:
#         for r in records:
#             txt = r["text"].replace("'", "''")
#             var = r["variation"]
#             valid = '1' if r.get("is_valid") else '0'
#             f.write(f"INSERT INTO sentences(text,variation,is_valid) VALUES('{txt}','{var}',{valid});\n")

# # -----------------------------------------------------------------------------
# # 10. Main
# # -----------------------------------------------------------------------------
# if __name__ == "__main__":
#     count   = int(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_COUNT
#     records = collect_records(count)
#     df      = pd.DataFrame(records)

#     emit_csv(df)
#     emit_xlsx(df)
#     emit_parquet(df)
#     emit_json(records)

#     emit_pst(records)

#     emit_txt(records)
#     emit_log(records)
#     emit_diff(records)
#     emit_html(records)

#     emit_paragraphs(records)

#     emit_pdf(records)
#     emit_docx(records)
#     emit_sql(records)

#     print(f"✅ Generated {count} records per module from {len(MODULES)} generators")
#     print(f"   → all outputs in folder: {OUT_DIR}")



#!/usr/bin/env python3
import importlib, json, os, random, sys
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

# -----------------------------------------------------------------------------
# 1. Configuration
# -----------------------------------------------------------------------------
DEFAULT_COUNT   = 50    # records per module
PARA_COUNT      = 5
SENTS_PER_PARA  = 5
OUT_DIR         = "mixed_output"
os.makedirs(OUT_DIR, exist_ok=True)

# -----------------------------------------------------------------------------
# 2. Discover & import your generate_*.py modules
# -----------------------------------------------------------------------------
MODULES = []
for fn in os.listdir("."):
    if fn.startswith("generate_") and fn.endswith(".py"):
        MODULES.append(importlib.import_module(fn[:-3]))

# -----------------------------------------------------------------------------
# 3. Find each module’s *_variations() function
# -----------------------------------------------------------------------------
def get_generator(mod):
    funcs = [getattr(mod, n) for n in dir(mod)
             if n.endswith("_variations") and callable(getattr(mod, n))]
    if not funcs:
        raise AttributeError(f"No `*_variations` in {mod.__name__}")
    if len(funcs)==1:
        return funcs[0]
    base = mod.__name__.split("generate_")[-1]
    for fn in funcs:
        if base in fn.__name__:
            return fn
    return funcs[0]

# -----------------------------------------------------------------------------
# 4. Collect & shuffle
# -----------------------------------------------------------------------------
def collect_records(count_per_module):
    pool = []
    for mod in MODULES:
        gen = get_generator(mod)
        pool.extend(gen(count_per_module))
    random.shuffle(pool)
    return pool

# -----------------------------------------------------------------------------
# 5. Structured output
# -----------------------------------------------------------------------------
def emit_structured(df):
    # CSV
    df.to_csv(os.path.join(OUT_DIR, "output.csv"), index=False)
    # filtered JSON (no nulls)
    records = df.to_dict(orient="records")
    cleaned = []
    for r in records:
        clean = {k: v for k, v in r.items() if pd.notna(v)}
        cleaned.append(clean)
    with open(os.path.join(OUT_DIR, "output.json"), "w", encoding="utf-8") as f:
        json.dump(cleaned, f, ensure_ascii=False, indent=2)
    # Excel .xlsx
    df.to_excel(os.path.join(OUT_DIR, "output.xlsx"), index=False)
    # Parquet
    try:
        df.to_parquet(os.path.join(OUT_DIR, "output.parquet"), index=False)
    except ImportError:
        pass

# -----------------------------------------------------------------------------
# 6. PST as SQL dump
# -----------------------------------------------------------------------------
def emit_pst(records):
    with open(os.path.join(OUT_DIR, "output.pst"), "w", encoding="utf-8") as f:
        for r in records:
            txt = r["text"].replace("'", "''")
            var = r["variation"]
            val = '1' if r.get("is_valid") else '0'
            f.write(
                f"INSERT INTO sentences(text,variation,is_valid)"
                f" VALUES('{txt}','{var}',{val});\n"
            )

# -----------------------------------------------------------------------------
# 7. Unstructured exports
# -----------------------------------------------------------------------------
def emit_txt(records):
    with open(os.path.join(OUT_DIR, "output.txt"), "w", encoding="utf-8") as f:
        for r in records:
            f.write(r["text"] + "\n\n")

def emit_log(records):
    with open(os.path.join(OUT_DIR, "output.log"), "w", encoding="utf-8") as f:
        for r in records:
            f.write(r["text"] + "\n")

def emit_html(records):
    with open(os.path.join(OUT_DIR, "output.html"), "w", encoding="utf-8") as f:
        f.write("<!doctype html>\n<html><body>\n")
        for r in records:
            # wrap each sentence in a <p>
            f.write(f"  <p>{r['text']}</p>\n")
        f.write("</body></html>")

def emit_pdf(records):
    c = canvas.Canvas(os.path.join(OUT_DIR, "output.pdf"), pagesize=letter)
    w, h = letter; y = h - 40
    for r in records:
        c.drawString(40, y, r["text"])
        y -= 14
        if y < 40:
            c.showPage(); y = h - 40
    c.save()

def emit_docx(records):
    doc = Document()
    for r in records:
        doc.add_paragraph(r["text"])
    doc.save(os.path.join(OUT_DIR, "output.docx"))

def emit_sql(records):
    with open(os.path.join(OUT_DIR, "output.sql"), "w", encoding="utf-8") as f:
        for r in records:
            txt = r["text"].replace("'", "''")
            var = r["variation"]
            val = '1' if r.get("is_valid") else '0'
            f.write(
                f"INSERT INTO sentences(text,variation,is_valid)"
                f" VALUES('{txt}','{var}',{val});\n"
            )

def emit_diff(records):
    with open(os.path.join(OUT_DIR, "output.diff"), "w", encoding="utf-8") as f:
        f.write("--- original\n+++ mixed\n")
        for r in records:
            f.write(f"+ {r['text']}\n")

def emit_paragraphs(records):
    needed = PARA_COUNT * SENTS_PER_PARA
    block  = records[:needed]
    with open(os.path.join(OUT_DIR, "output_paras.txt"), "w", encoding="utf-8") as f:
        for i in range(PARA_COUNT):
            chunk = block[i*SENTS_PER_PARA:(i+1)*SENTS_PER_PARA]
            para = " ".join(r["text"] for r in chunk)
            f.write(para + "\n\n")

# -----------------------------------------------------------------------------
# 8. Main
# -----------------------------------------------------------------------------
if __name__ == "__main__":
    count   = int(sys.argv[1]) if len(sys.argv)>1 else DEFAULT_COUNT
    records = collect_records(count)
    df      = pd.DataFrame(records)

    emit_structured(df)
    emit_pst(records)          # now SQL, not JSON
    emit_txt(records)
    emit_log(records)
    emit_html(records)         # real HTML
    emit_pdf(records)
    emit_docx(records)
    emit_sql(records)          # INSERT statements
    emit_diff(records)
    emit_paragraphs(records)

    print(f"✅ Generated {count} records per module from {len(MODULES)} generators")
    print(f"   → all outputs in folder: {OUT_DIR}")
